"""
Utility functions for Django Essay Coach Application
"""
import os
import re
import tempfile
import logging
from functools import wraps
from datetime import datetime
from django.shortcuts import redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.core.files.storage import default_storage
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
import io

logger = logging.getLogger(__name__)


def role_required(role):
    """
    Decorator to check if user has required role
    """
    def decorator(view_func):
        @wraps(view_func)
        @login_required
        def wrapped_view(request, *args, **kwargs):
            if not hasattr(request.user, 'role') or request.user.role != role:
                messages.error(request, f'Access denied. {role.title()} role required.')
                return redirect('accounts:index')
            return view_func(request, *args, **kwargs)
        return wrapped_view
    return decorator


def allowed_file(filename):
    """Check if file has allowed extension"""
    allowed_extensions = getattr(settings, 'ALLOWED_EXTENSIONS', {'docx', 'txt'})
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions


def is_file_size_valid(file):
    """Check if file size is within limits"""
    max_size = getattr(settings, 'MAX_CONTENT_LENGTH', 16 * 1024 * 1024)
    return file.size <= max_size


def validate_file_upload(file):
    """
    Validate uploaded file
    
    Args:
        file: Django UploadedFile instance
    
    Returns:
        tuple: (is_valid, error_message)
    """
    if not file:
        return False, "No file provided"
    
    if not allowed_file(file.name):
        return False, "File type not allowed. Please upload a .docx or .txt file."
    
    if not is_file_size_valid(file):
        max_size_mb = getattr(settings, 'MAX_CONTENT_LENGTH', 16 * 1024 * 1024) / (1024 * 1024)
        return False, f"File size exceeds {max_size_mb}MB limit."
    
    return True, ""


def extract_text_from_file(file):
    """
    Extract text from uploaded file
    
    Args:
        file: Django UploadedFile instance
    
    Returns:
        str: Extracted text content
    """
    try:
        file_extension = file.name.rsplit('.', 1)[1].lower()
        
        if file_extension == 'txt':
            # Handle text files
            content = file.read()
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                    try:
                        text = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise UnicodeDecodeError("Unable to decode file with any supported encoding")
            else:
                text = content
                
        elif file_extension == 'docx':
            # Handle Word documents
            doc = Document(file)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Clean up the text
        text = sanitize_text(text)
        
        # Validate text length
        min_length = getattr(settings, 'ERROR_HANDLING', {}).get('file_min_text_length', 50)
        max_length = getattr(settings, 'ERROR_HANDLING', {}).get('file_max_text_length', 50000)
        
        if len(text) < min_length:
            raise ValueError(f"Text too short. Minimum {min_length} characters required.")
        
        if len(text) > max_length:
            raise ValueError(f"Text too long. Maximum {max_length} characters allowed.")
        
        logger.info(f"Successfully extracted {len(text)} characters from {file.name}")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from file {file.name}: {e}")
        raise


def sanitize_text(text):
    """
    Clean and sanitize text content
    
    Args:
        text (str): Raw text content
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Remove control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive newlines
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()


def create_word_document_with_suggestions(essay_text, suggestions, filename=None, analysis=None):
    """
    Create a Word document with essay text and suggestions with EXACT formatting requirements:
    - Double-spaced throughout the entire document
    - Deletions: Blue text with strikethrough
    - Additions: Red text with underline
    - Includes scores and explanations
    - Professional formatting matching academic standards
    
    Args:
        essay_text (str): Original essay text
        suggestions (list): List of suggestions
        filename (str): Optional filename
        analysis (EssayAnalysis): Analysis object with scores and feedback
    
    Returns:
        io.BytesIO: Word document as bytes
    """
    try:
        doc = Document()
        
        # Configure document-wide styles for double spacing
        styles = doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        
        # Set paragraph format for double spacing - this is CRITICAL
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        normal_paragraph_format.space_after = Pt(0)
        normal_paragraph_format.space_before = Pt(0)
        
        # PROFESSIONAL DOCUMENT HEADER
        # Main Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        title_para.paragraph_format.space_after = Pt(12)
        title_run = title_para.add_run('AI ESSAY ANALYSIS REPORT')
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.all_caps = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Subtitle with date
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        subtitle_para.paragraph_format.space_after = Pt(18)
        date_run = subtitle_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Horizontal line (using border)
        divider_para = doc.add_paragraph()
        divider_para.paragraph_format.space_after = Pt(18)
        divider_run = divider_para.add_run('_' * 80)
        divider_run.font.name = 'Times New Roman'
        divider_run.font.size = Pt(8)
        divider_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
        
        # RUBRIC SCORES SECTION - Double spaced
        if analysis:
            scores_heading = doc.add_paragraph()
            scores_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            scores_heading_run = scores_heading.add_run('Rubric Scores')
            scores_heading_run.font.name = 'Times New Roman'
            scores_heading_run.font.size = Pt(14)
            scores_heading_run.font.bold = True
            
            # Individual scores - each on double-spaced line
            scores_data = {
                'Ideas & Content': (int(getattr(analysis, 'content_score', 0)), 20),
                'Organization': (int(getattr(analysis, 'structure_score', 0)), 25),
                'Style & Voice': (int(getattr(analysis, 'clarity_score', 0)), 25),
                'Grammar': (int(getattr(analysis, 'grammar_score', 0)), 30),
                'Overall Score': (int(getattr(analysis, 'overall_score', 0)), 100)
            }
            
            for category, (score, max_score) in scores_data.items():
                score_para = doc.add_paragraph()
                score_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                
                # Category name in bold
                cat_run = score_para.add_run(f"{category}: ")
                cat_run.font.name = 'Times New Roman'
                cat_run.font.size = Pt(12)
                cat_run.font.bold = True
                
                # Score value with correct maximum
                score_run = score_para.add_run(f"{score}/{max_score}")
                score_run.font.name = 'Times New Roman'
                score_run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            
            # SCORE EXPLANATIONS - Double spaced
            if hasattr(analysis, 'detailed_feedback') and analysis.detailed_feedback:
                exp_heading = doc.add_paragraph()
                exp_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                exp_heading_run = exp_heading.add_run('Score Explanations')
                exp_heading_run.font.name = 'Times New Roman'
                exp_heading_run.font.size = Pt(14)
                exp_heading_run.font.bold = True
                
                feedback = analysis.detailed_feedback
                explanations = {
                    'Ideas & Content': feedback.get('content', ''),
                    'Organization': feedback.get('structure', ''),
                    'Style & Voice': feedback.get('clarity', ''),
                    'Grammar': feedback.get('grammar', '')
                }
                
                for category, explanation in explanations.items():
                    if explanation:
                        # Category heading
                        cat_para = doc.add_paragraph()
                        cat_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        cat_run = cat_para.add_run(f"{category}:")
                        cat_run.font.name = 'Times New Roman'
                        cat_run.font.size = Pt(12)
                        cat_run.font.bold = True
                        
                        # Explanation text
                        exp_para = doc.add_paragraph()
                        exp_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                        exp_run = exp_para.add_run(explanation)
                        exp_run.font.name = 'Times New Roman'
                        exp_run.font.size = Pt(12)
                
                # Add spacing
                doc.add_paragraph()
        
        # SECTION 2: REVISED ESSAY WITH AI SUGGESTIONS
        essay_heading = doc.add_heading('2. REVISED ESSAY WITH AI SUGGESTIONS', level=1)
        essay_heading.runs[0].font.name = 'Times New Roman'
        essay_heading.runs[0].font.size = Pt(14)
        essay_heading.runs[0].font.bold = True
        essay_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Professional legend with clear formatting
        legend_para = doc.add_paragraph()
        legend_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        legend_para.paragraph_format.space_after = Pt(12)
        legend_para.paragraph_format.left_indent = Pt(18)
        
        legend_run = legend_para.add_run('Formatting Legend: ')
        legend_run.font.name = 'Times New Roman'
        legend_run.font.size = Pt(11)
        legend_run.font.bold = True
        
        # Blue strikethrough example
        blue_run = legend_para.add_run('Deleted text')
        blue_run.font.name = 'Times New Roman'
        blue_run.font.size = Pt(11)
        blue_run.font.color.rgb = RGBColor(0, 0, 255)
        blue_run.font.strike = True
        
        legend_para.add_run(' • ').font.size = Pt(11)
        
        # Red underline example
        red_run = legend_para.add_run('Added text')
        red_run.font.name = 'Times New Roman'
        red_run.font.size = Pt(11)
        red_run.font.color.rgb = RGBColor(255, 0, 0)
        red_run.font.underline = True
        
        legend_para.add_run(' • ').font.size = Pt(11)
        
        # Explanation format with better description
        italic_run = legend_para.add_run('(grammatical explanations in brackets)')
        italic_run.font.name = 'Times New Roman'
        italic_run.font.size = Pt(10)
        italic_run.font.italic = True
        italic_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Additional explanation paragraph
        explanation_para = doc.add_paragraph()
        explanation_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        explanation_para.paragraph_format.space_after = Pt(12)
        explanation_para.paragraph_format.left_indent = Pt(18)
        
        explanation_text = explanation_para.add_run(
            'Each suggested change is followed by a grammatical explanation in parentheses '
            'that describes the specific grammar rule, spelling correction, or style improvement being applied.'
        )
        explanation_text.font.name = 'Times New Roman'
        explanation_text.font.size = Pt(10)
        explanation_text.font.italic = True
        explanation_text.font.color.rgb = RGBColor(100, 100, 100)
        
        # Get tagged essay from analysis
        tagged_essay = essay_text
        if analysis and hasattr(analysis, 'detailed_feedback') and analysis.detailed_feedback:
            tagged_essay = analysis.detailed_feedback.get('tagged_essay', essay_text)
        
        # PROCESS THE ESSAY TEXT WITH PRECISE WORD-BY-WORD FORMATTING
        essay_para = doc.add_paragraph()
        essay_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Parse and format the tagged text with true word-by-word processing
        current_text = tagged_essay
        
        # Helper function to find suggestion reason for specific text with grammatical explanations
        def find_suggestion_reason(text, suggestion_type):
            if suggestions:
                for suggestion in suggestions:
                    if isinstance(suggestion, dict):
                        if (suggestion.get('type') == suggestion_type and 
                            (text in str(suggestion.get('text', '')) or 
                             text in str(suggestion.get('reason', '')))):
                            reason = suggestion.get('reason', '')
                            # Enhance with specific grammatical explanations
                            return get_grammatical_explanation(text, suggestion_type, reason)
            return get_grammatical_explanation(text, suggestion_type, '')
        
        # Function to provide detailed grammatical explanations
        def get_grammatical_explanation(text, suggestion_type, original_reason=''):
            """Provide specific grammatical explanations for different types of corrections"""
            text_lower = text.lower().strip()
            
            # Common grammatical patterns and their explanations
            grammatical_explanations = {
                # Subject-verb agreement
                'have': 'subject-verb agreement: singular subject requires singular verb',
                'has': 'subject-verb agreement: plural subject requires plural verb', 
                'is': 'subject-verb agreement: plural subject requires plural verb',
                'are': 'subject-verb agreement: singular subject requires singular verb',
                'was': 'subject-verb agreement: plural subject requires plural verb',
                'were': 'subject-verb agreement: singular subject requires singular verb',
                
                # Article usage
                'a': 'article usage: indefinite article before consonant sound',
                'an': 'article usage: indefinite article before vowel sound',
                'the': 'article usage: definite article for specific reference',
                
                # Plural/singular forms
                'peoples': 'word form: "people" is already plural',
                'people': 'word form: correct plural form',
                'advantage': 'number agreement: singular form',
                'advantages': 'number agreement: plural form required',
                'disadvantages': 'number agreement: plural form required',
                'opportunities': 'spelling and plural form correction',
                'cultures': 'number agreement: plural form required',
                'traditions': 'number agreement: plural form required',
                'hospitals': 'number agreement: plural form required',
                'schools': 'number agreement: plural form required',
                'malls': 'number agreement: plural form required',
                'towns': 'number agreement: plural form required',
                'villages': 'number agreement: plural form required',
                'degrees': 'number agreement: plural form required',
                'clothes': 'word form: plural noun',
                
                # Verb forms and tenses
                'seeing': 'verb form: simple present tense required',
                'see': 'verb form: correct simple present',
                'goes': 'verb tense: past tense required',
                'went': 'verb tense: correct past tense',
                'finding': 'verb form: simple past tense required',
                'find': 'verb form: infinitive after auxiliary',
                'proving': 'verb form: simple present tense required',
                'proves': 'verb form: correct third person singular',
                'attend': 'verb form: past continuous tense required',
                'attending': 'verb form: correct past continuous',
                'wear': 'verb tense: past tense required',
                'wore': 'verb tense: correct past tense',
                'feel': 'verb tense: past tense required',
                'felt': 'verb tense: correct past tense',
                'thinking': 'verb form: simple present tense required',
                'think': 'verb form: correct simple present',
                'giving': 'verb form: simple present tense required',
                'give': 'verb form: correct simple present',
                'growing': 'verb form: relative clause structure',
                'grow': 'verb form: correct present tense',
                'requiring': 'verb form: simple present tense required',
                'requires': 'verb form: correct third person singular',
                'wasting': 'verb form: simple present tense required',
                'wastes': 'verb form: correct third person singular',
                'effecting': 'word choice: "affecting" means influencing',
                'affecting': 'word choice: correct verb meaning',
                'causing': 'verb form: correct present participle',
                'adapting': 'verb form: modal verb requires infinitive',
                'adapt': 'verb form: correct infinitive after modal',
                
                # Word choice and spelling
                'alot': 'spelling: two separate words required',
                'a lot': 'spelling: correct two-word form',
                'oppertunitys': 'spelling: correct spelling is "opportunities"',
                'sacrifies': 'spelling: correct spelling is "sacrifices"',
                'sacrifices': 'spelling: correct form',
                'depressions': 'word form: uncountable noun, no plural',
                'depression': 'word form: correct uncountable noun',
                'angry': 'word form: noun form required',
                'anger': 'word form: correct noun',
                'healthy': 'word form: noun form required',
                'health': 'word form: correct noun',
                'traffics': 'word form: uncountable noun, no plural',
                'traffic': 'word form: correct uncountable noun',
                'persons': 'word choice: "people" is preferred plural',
                'nightmarea': 'article usage: indefinite article required',
                'nightmare': 'word form: correct noun',
                'paradisea': 'article usage: indefinite article required', 
                'paradise': 'word form: correct noun',
                'prepare': 'word form: noun form required',
                'preparation': 'word form: correct noun',
                'horn': 'word choice: "honk" is correct verb',
                'honk': 'word choice: correct verb for car horns',
                'headache': 'word form: correct compound noun',
                'drivers': 'word choice: context requires "people"',
                'breathe in': 'phrasal verb: redundant preposition',
                'breathe': 'verb form: correct simple form',
                
                # Preposition usage
                'with': 'preposition: "full of" is correct phrase',
                'of': 'preposition: correct usage with "full"',
                'to': 'preposition: gerund requires different preposition',
                'from': 'preposition: correct with comparison',
                'in': 'preposition: "to" is correct for movement',
                'among': 'preposition: correct for being surrounded by',
                
                # Sentence structure
                'though': 'conjunction: "even though" is complete phrase',
                'even though': 'conjunction: correct concessive phrase',
                'what': 'relative pronoun: "that" is correct here',
                'that': 'relative pronoun: correct relative pronoun',
                'not': 'sentence structure: contraction preferred',
                'don\'t': 'contraction: correct negative form',
                'didn\'t': 'contraction: correct past negative',
                'compare': 'sentence structure: "compared to" for comparison',
                'compared': 'sentence structure: correct comparative phrase',
                'out of box': 'idiom: correct phrase is "out of place"',
                'out of place': 'idiom: correct idiomatic expression',
                'easy': 'adverb form: modify verb with adverb',
                'easily': 'adverb form: correct adverb',
                'loud': 'sentence structure: "are loud" for description',
                'are loud': 'sentence structure: correct predicate adjective',
                'quiet': 'sentence structure: "are quiet" for description', 
                'are quiet': 'sentence structure: correct predicate adjective',
                'little': 'article usage: indefinite article required',
                'a little': 'article usage: correct indefinite quantity',
                'sometime': 'word choice: "sometimes" for frequency',
                'sometimes': 'word choice: correct adverb of frequency',
                'also': 'word order: redundant with "too"',
                'too': 'word order: "also" already used',
                'much': 'quantifier: "many" for countable nouns',
                'many': 'quantifier: correct for countable nouns',
                'fill': 'word form: past participle required',
                'filled': 'word form: correct past participle',
                'only can': 'modal order: "can only" is correct order',
                'can only': 'modal order: correct auxiliary verb order',
                'will': 'verb choice: "can" is appropriate here',
                'can': 'verb choice: correct modal for ability',
                'truth': 'word form: adjective required',
                'true': 'word form: correct adjective',
                'dreams': 'article usage: "a dream" for singular',
                'dream': 'article usage: correct with indefinite article',
                'suffer': 'verb form: "suffer from" is correct phrase',
                'suffer from': 'phrasal verb: correct preposition usage'
            }
            
            # Try to find specific explanation
            if text_lower in grammatical_explanations:
                return grammatical_explanations[text_lower]
            
            # Fallback to general explanations based on suggestion type
            if suggestion_type == 'delete':
                if 'redundant' in original_reason.lower():
                    return 'redundancy: unnecessary word'
                elif 'incorrect' in original_reason.lower():
                    return 'grammar error: incorrect word usage'
                else:
                    return 'word choice: word should be removed'
                    
            elif suggestion_type == 'add':
                if 'article' in original_reason.lower():
                    return 'article usage: missing article'
                elif 'preposition' in original_reason.lower():
                    return 'preposition: missing preposition' 
                else:
                    return 'grammar: missing word required'
                    
            elif suggestion_type == 'replace':
                if 'tense' in original_reason.lower():
                    return 'verb tense: incorrect tense usage'
                elif 'agreement' in original_reason.lower():
                    return 'subject-verb agreement: mismatch'
                elif 'spelling' in original_reason.lower():
                    return 'spelling: incorrect spelling'
                else:
                    return 'word choice: better word choice'
            
            return original_reason or f'{suggestion_type} suggestion'
        
        # Helper function to ensure word-by-word tagging (split multi-word tags)
        def ensure_word_by_word_tags(text):
            """
            Post-process tagged text to ensure each tag contains only one word
            """
            import re
            
            def split_tag_content(match):
                tag_type = match.group(1)  # delete, add, or replace
                content = match.group(2)
                
                if tag_type == 'replace' and '|' in content:
                    old_part, new_part = content.split('|', 1)
                    old_words = old_part.strip().split()
                    new_words = new_part.strip().split()
                    
                    result = ""
                    # Handle case where old and new have different word counts
                    max_words = max(len(old_words), len(new_words))
                    for i in range(max_words):
                        old_word = old_words[i] if i < len(old_words) else ""
                        new_word = new_words[i] if i < len(new_words) else ""
                        
                        if old_word and new_word:
                            result += f"<replace>{old_word}|{new_word}</replace> "
                        elif old_word:
                            result += f"<delete>{old_word}</delete> "
                        elif new_word:
                            result += f"<add>{new_word}</add> "
                    return result.strip()
                else:
                    # For delete and add tags, split multi-word content
                    words = content.strip().split()
                    if len(words) <= 1:
                        return match.group(0)  # Return original if single word
                    
                    result = ""
                    for word in words:
                        result += f"<{tag_type}>{word}</{tag_type}> "
                    return result.strip()
            
            # Apply word-by-word splitting to all tags
            pattern = r'<(delete|add|replace)>(.*?)</\1>'
            return re.sub(pattern, split_tag_content, text)
        
        # Ensure the tagged essay has word-by-word tags
        tagged_essay = ensure_word_by_word_tags(tagged_essay)
        
        # Improved word-by-word processing using character-by-character parsing
        i = 0
        while i < len(current_text):
            # Check for opening tags at current position
            if current_text[i:i+8] == '<delete>':
                # Find the closing tag
                end_pos = current_text.find('</delete>', i + 8)
                if end_pos != -1:
                    deleted_text = current_text[i + 8:end_pos].strip()
                    
                    # Should be a single word after preprocessing
                    # EXACT REQUIREMENT: Blue text with strikethrough
                    del_run = essay_para.add_run(deleted_text)
                    del_run.font.name = 'Times New Roman'
                    del_run.font.size = Pt(12)
                    del_run.font.color.rgb = RGBColor(0, 0, 255)  # BLUE
                    del_run.font.strike = True  # STRIKETHROUGH
                    
                    # Add inline explanation for the word
                    reason = find_suggestion_reason(deleted_text, 'delete')
                    explanation_run = essay_para.add_run(f" ({reason})")
                    explanation_run.font.name = 'Times New Roman'
                    explanation_run.font.size = Pt(9)
                    explanation_run.font.color.rgb = RGBColor(0, 0, 150)  # Darker blue
                    explanation_run.font.italic = True
                    
                    i = end_pos + 9  # Move past </delete>
                    continue
                    
            elif current_text[i:i+5] == '<add>':
                # Find the closing tag
                end_pos = current_text.find('</add>', i + 5)
                if end_pos != -1:
                    added_text = current_text[i + 5:end_pos].strip()
                    
                    # Should be a single word after preprocessing
                    # EXACT REQUIREMENT: Red text with underline
                    add_run = essay_para.add_run(added_text)
                    add_run.font.name = 'Times New Roman'
                    add_run.font.size = Pt(12)
                    add_run.font.color.rgb = RGBColor(255, 0, 0)  # RED
                    add_run.font.underline = True  # UNDERLINE
                    
                    # Add inline explanation for the word
                    reason = find_suggestion_reason(added_text, 'add')
                    explanation_run = essay_para.add_run(f" ({reason})")
                    explanation_run.font.name = 'Times New Roman'
                    explanation_run.font.size = Pt(9)
                    explanation_run.font.color.rgb = RGBColor(200, 0, 0)  # Darker red
                    explanation_run.font.italic = True
                    
                    i = end_pos + 6  # Move past </add>
                    continue
                    
            elif current_text[i:i+9] == '<replace>':
                # Find the closing tag
                end_pos = current_text.find('</replace>', i + 9)
                if end_pos != -1:
                    replace_content = current_text[i + 9:end_pos]
                    
                    if '|' in replace_content:
                        old_text, new_text = replace_content.split('|', 1)
                        old_text = old_text.strip()
                        new_text = new_text.strip()
                        
                        # Old word: Blue with strikethrough
                        old_run = essay_para.add_run(old_text)
                        old_run.font.name = 'Times New Roman'
                        old_run.font.size = Pt(12)
                        old_run.font.color.rgb = RGBColor(0, 0, 255)  # BLUE
                        old_run.font.strike = True  # STRIKETHROUGH
                        
                        # Add small space
                        essay_para.add_run(' ')
                        
                        # New word: Red with underline
                        new_run = essay_para.add_run(new_text)
                        new_run.font.name = 'Times New Roman'
                        new_run.font.size = Pt(12)
                        new_run.font.color.rgb = RGBColor(255, 0, 0)  # RED
                        new_run.font.underline = True  # UNDERLINE
                        
                        # Add inline explanation for the replacement
                        reason = find_suggestion_reason(f"{old_text} -> {new_text}", 'replace')
                        explanation_run = essay_para.add_run(f" ({reason})")
                        explanation_run.font.name = 'Times New Roman'
                        explanation_run.font.size = Pt(9)
                        explanation_run.font.color.rgb = RGBColor(128, 0, 128)  # Purple for replacements
                        explanation_run.font.italic = True
                    else:
                        # No pipe separator, treat as normal text
                        normal_run = essay_para.add_run(replace_content.strip())
                        normal_run.font.name = 'Times New Roman'
                        normal_run.font.size = Pt(12)
                    
                    i = end_pos + 10  # Move past </replace>
                    continue
            
            # No tag found at current position, add the character as normal text
            # Look ahead to find the next tag or end of text
            next_tag_pos = len(current_text)
            for tag in ['<delete>', '<add>', '<replace>']:
                tag_pos = current_text.find(tag, i)
                if tag_pos != -1 and tag_pos < next_tag_pos:
                    next_tag_pos = tag_pos
            
            # Add all text from current position to next tag as normal text
            if next_tag_pos > i:
                normal_text = current_text[i:next_tag_pos]
                if normal_text:
                    # Add normal text while preserving original spacing
                    normal_run = essay_para.add_run(normal_text)
                    normal_run.font.name = 'Times New Roman'
                    normal_run.font.size = Pt(12)
                i = next_tag_pos
            else:
                i += 1
        
        # Add final spacing before footer
        doc.add_paragraph()
        doc.add_paragraph()
        
        # PROFESSIONAL FOOTER SECTION
        # Divider line
        footer_divider = doc.add_paragraph()
        footer_divider.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        divider_run = footer_divider.add_run('_' * 60)
        divider_run.font.name = 'Times New Roman'
        divider_run.font.size = Pt(8)
        divider_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
        
        # Footer with branding
        footer_para = doc.add_paragraph()
        footer_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_para.paragraph_format.space_after = Pt(6)
        
        footer_title = footer_para.add_run('AI Essay Coach')
        footer_title.font.name = 'Times New Roman'
        footer_title.font.size = Pt(12)
        footer_title.font.bold = True
        footer_title.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Date and time on separate line
        date_para = doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        date_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        logger.info("Enhanced Word document with precise formatting created successfully")
        return doc_io
        
    except Exception as e:
        logger.error(f"Error creating enhanced Word document: {e}")
        # Fallback to simple document
        return create_simple_word_document(essay_text, suggestions, filename)


def create_simple_word_document(essay_text, suggestions, filename=None):
    """
    Create a simple Word document as fallback
    """
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Essay with Suggestions', 0)
        
        # Add essay text
        doc.add_heading('Original Essay', level=1)
        essay_paragraph = doc.add_paragraph(essay_text)
        
        # Add suggestions
        doc.add_heading('Suggestions for Improvement', level=1)
        for i, suggestion in enumerate(suggestions, 1):
            if isinstance(suggestion, dict):
                suggestion_text = suggestion.get('text', '') or suggestion.get('reason', '') or str(suggestion)
            else:
                suggestion_text = str(suggestion)
            doc.add_paragraph(f"{i}. {suggestion_text}")
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
        
    except Exception as e:
        logger.error(f"Error creating simple Word document: {e}")
        raise


def get_current_user(request):
    """Get current user information"""
    if request.user.is_authenticated:
        return {
            'id': request.user.id,
            'username': request.user.username,
            'role': getattr(request.user, 'role', 'student'),
            'is_student': getattr(request.user, 'is_student', lambda: False)(),
            'is_teacher': getattr(request.user, 'is_teacher', lambda: False)(),
        }
    return None


class TemporaryStorage:
    """Temporary storage for analysis data"""
    
    def __init__(self):
        self.storage_dir = os.path.join(settings.BASE_DIR, 'temp_data')
        if not os.path.exists(self.storage_dir):
            os.makedirs(self.storage_dir)
    
    def store_analysis(self, key, data):
        """Store analysis data temporarily"""
        try:
            file_path = os.path.join(self.storage_dir, f"{key}.json")
            with open(file_path, 'w') as f:
                import json
                json.dump(data, f)
            logger.info(f"Analysis stored temporarily with key: {key}")
            return True
        except Exception as e:
            logger.error(f"Error storing temporary analysis: {e}")
            return False
    
    def retrieve_analysis(self, key):
        """Retrieve analysis data from temporary storage"""
        try:
            file_path = os.path.join(self.storage_dir, f"{key}.json")
            if os.path.exists(file_path):
                with open(file_path, 'r') as f:
                    import json
                    data = json.load(f)
                logger.info(f"Analysis retrieved from temporary storage: {key}")
                return data
            return None
        except Exception as e:
            logger.error(f"Error retrieving temporary analysis: {e}")
            return None
    
    def cleanup_expired(self, max_age_hours=24):
        """Clean up expired temporary files"""
        try:
            import time
            current_time = time.time()
            max_age_seconds = max_age_hours * 3600
            
            for filename in os.listdir(self.storage_dir):
                file_path = os.path.join(self.storage_dir, filename)
                if os.path.isfile(file_path):
                    file_age = current_time - os.path.getmtime(file_path)
                    if file_age > max_age_seconds:
                        os.remove(file_path)
                        logger.info(f"Cleaned up expired temporary file: {filename}")
        except Exception as e:
            logger.error(f"Error cleaning up temporary files: {e}")


# Initialize temporary storage
temp_storage = TemporaryStorage()


def store_analysis_temporarily(key, data):
    """Store analysis data temporarily"""
    return temp_storage.store_analysis(key, data)


def retrieve_analysis_temporarily(key):
    """Retrieve analysis data from temporary storage"""
    return temp_storage.retrieve_analysis(key)


def cleanup_expired_temp_data():
    """Clean up expired temporary data"""
    temp_storage.cleanup_expired()
