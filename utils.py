"""
Utilities module for the Essay Revision Application
Contains utility functions for file handling, validation, and document processing
"""
import os
import re
import logging
import json
import uuid
import time
from flask import session, redirect, url_for, flash
from functools import wraps
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_LINE_SPACING, WD_UNDERLINE
import datetime
from io import BytesIO
from config import Config

logger = logging.getLogger(__name__)

class TemporaryDataStorage:
    """
    Temporary file-based storage for large data that shouldn't be stored in session
    """
    
    def __init__(self, storage_dir: str = None, ttl: int = None):
        """
        Initialize temporary storage
        
        Args:
            storage_dir (str): Directory for temporary files (default from config)
            ttl (int): Time to live in seconds (default from config)
        """
        self.storage_dir = storage_dir or Config.TEMP_STORAGE.get('directory', 'temp_data')
        self.ttl = ttl or Config.TEMP_STORAGE.get('ttl', 3600)
        
        # Make path absolute
        if not os.path.isabs(self.storage_dir):
            self.storage_dir = os.path.join(os.getcwd(), self.storage_dir)
        
        # Ensure storage directory exists
        os.makedirs(self.storage_dir, exist_ok=True)
    
    def store(self, data: dict) -> str:
        """
        Store data temporarily and return a unique identifier
        
        Args:
            data (dict): Data to store
            
        Returns:
            str: Unique identifier for the stored data
        """
        try:
            # Generate unique identifier
            data_id = str(uuid.uuid4())
            
            # Add timestamp for TTL
            storage_data = {
                'data': data,
                'timestamp': time.time(),
                'ttl': self.ttl
            }
            
            # Store in file
            file_path = os.path.join(self.storage_dir, f"{data_id}.json")
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(storage_data, f, default=str, ensure_ascii=False, indent=2)
            
            logger.info(f"Temporary data stored with ID: {data_id}")
            return data_id
            
        except Exception as e:
            logger.error(f"Error storing temporary data: {e}")
            return None
    
    def retrieve(self, data_id: str, delete_after_read: bool = True) -> dict:
        """
        Retrieve temporarily stored data
        
        Args:
            data_id (str): Unique identifier
            delete_after_read (bool): Whether to delete the file after reading
            
        Returns:
            dict: Stored data or None if not found/expired
        """
        try:
            file_path = os.path.join(self.storage_dir, f"{data_id}.json")
            
            if not os.path.exists(file_path):
                logger.warning(f"Temporary data file not found: {data_id}")
                return None
            
            with open(file_path, 'r', encoding='utf-8') as f:
                storage_data = json.load(f)
            
            # Check if data has expired
            if time.time() - storage_data['timestamp'] > storage_data['ttl']:
                logger.warning(f"Temporary data expired: {data_id}")
                self._delete_file(file_path)
                return None
            
            # Delete file if requested
            if delete_after_read:
                self._delete_file(file_path)
                logger.info(f"Temporary data retrieved and deleted: {data_id}")
            else:
                logger.info(f"Temporary data retrieved: {data_id}")
            
            return storage_data['data']
            
        except Exception as e:
            logger.error(f"Error retrieving temporary data {data_id}: {e}")
            return None
    
    def cleanup_expired(self):
        """Remove expired temporary files"""
        try:
            current_time = time.time()
            cleaned_count = 0
            
            for filename in os.listdir(self.storage_dir):
                if filename.endswith('.json'):
                    file_path = os.path.join(self.storage_dir, filename)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            storage_data = json.load(f)
                        
                        # Check if expired
                        if current_time - storage_data['timestamp'] > storage_data['ttl']:
                            self._delete_file(file_path)
                            cleaned_count += 1
                            
                    except Exception as e:
                        logger.warning(f"Error checking expiry for {filename}: {e}")
                        # Delete corrupted files
                        self._delete_file(file_path)
                        cleaned_count += 1
            
            if cleaned_count > 0:
                logger.info(f"Cleaned up {cleaned_count} expired temporary files")
                
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")
    
    def _delete_file(self, file_path: str):
        """Safely delete a file"""
        try:
            os.remove(file_path)
        except Exception as e:
            logger.warning(f"Could not delete temporary file {file_path}: {e}")

# Global temporary storage instance
temp_storage = TemporaryDataStorage()

def store_analysis_temporarily(analysis_data: dict) -> str:
    """
    Store analysis data temporarily and return an identifier
    
    Args:
        analysis_data (dict): Analysis data containing essay_text, title, and analysis
        
    Returns:
        str: Temporary storage identifier
    """
    return temp_storage.store(analysis_data)

def retrieve_analysis_temporarily(data_id: str) -> dict:
    """
    Retrieve temporarily stored analysis data
    
    Args:
        data_id (str): Temporary storage identifier
        
    Returns:
        dict: Analysis data or None if not found
    """
    return temp_storage.retrieve(data_id, delete_after_read=True)

def cleanup_expired_temp_data():
    """Clean up expired temporary data files"""
    temp_storage.cleanup_expired()

def schedule_cleanup():
    """Schedule periodic cleanup of temporary files"""
    import threading
    import time
    
    def cleanup_worker():
        while True:
            try:
                time.sleep(3600)  # Run every hour
                cleanup_expired_temp_data()
            except Exception as e:
                logger.error(f"Error in cleanup worker: {e}")
    
    # Run cleanup in background thread
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()
    logger.info("Temporary file cleanup scheduler started")

# Start cleanup scheduler when module is imported
schedule_cleanup()

class FileStreamer:
    """
    File streaming utility for handling large file uploads efficiently
    """
    
    def __init__(self, chunk_size: int = None, memory_threshold: int = None):
        """
        Initialize FileStreamer
        
        Args:
            chunk_size (int): Size of chunks for streaming (default from config)
            memory_threshold (int): Threshold for switching to streaming (default from config)
        """
        self.chunk_size = chunk_size or Config.PERFORMANCE.get('file_chunk_size', 8192)
        self.memory_threshold = memory_threshold or Config.PERFORMANCE.get('file_memory_threshold', 1024 * 1024)
    
    def should_stream(self, file_size: int) -> bool:
        """
        Determine if a file should be streamed based on size
        
        Args:
            file_size (int): Size of the file in bytes
            
        Returns:
            bool: True if file should be streamed
        """
        return (Config.PERFORMANCE.get('file_streaming_enabled', True) and 
                file_size > self.memory_threshold)
    
    def save_uploaded_file_streaming(self, file_storage, save_path: str) -> bool:
        """
        Save uploaded file using streaming for large files
        
        Args:
            file_storage: FileStorage object from form upload
            save_path (str): Path where to save the file
            
        Returns:
            bool: True if saved successfully
        """
        try:
            # Create directory if it doesn't exist
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            # Get file size if possible
            file_size = 0
            try:
                file_storage.seek(0, 2)  # Seek to end
                file_size = file_storage.tell()
                file_storage.seek(0)  # Reset to beginning
            except Exception:
                logger.warning("Could not determine file size for streaming decision")
            
            if self.should_stream(file_size):
                logger.info(f"Streaming upload of large file ({file_size} bytes) to {save_path}")
                with open(save_path, 'wb') as f:
                    for chunk in self.stream_file_chunks(file_storage):
                        f.write(chunk)
            else:
                logger.info(f"Saving small file ({file_size} bytes) to {save_path}")
                file_storage.save(save_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Error saving uploaded file to {save_path}: {e}")
            return False
    
    def stream_file_chunks(self, file_obj):
        """
        Stream file content in chunks
        
        Args:
            file_obj: File object to stream
            
        Yields:
            bytes: File chunks
        """
        try:
            while True:
                chunk = file_obj.read(self.chunk_size)
                if not chunk:
                    break
                
                # Ensure we're yielding bytes
                if isinstance(chunk, str):
                    chunk = chunk.encode('utf-8')
                
                yield chunk
        except Exception as e:
            logger.error(f"Error streaming file: {e}")
            raise
    
    def extract_text_from_file_streaming(self, file_path: str, file_type: str = None) -> str:
        """
        Extract text from file using streaming for large files
        
        Args:
            file_path (str): Path to the file
            file_type (str): Optional file type hint
            
        Returns:
            str: Extracted text content
        """
        try:
            if not file_type:
                file_type = file_path.rsplit('.', 1)[1].lower() if '.' in file_path else ''
            
            if file_type == 'txt':
                return self._extract_text_streaming(file_path)
            elif file_type == 'docx':
                return self._extract_docx_streaming(file_path)
            else:
                logger.warning(f"Unsupported file type for streaming: {file_type}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_text_streaming(self, file_path: str) -> str:
        """Extract text from TXT file with streaming support"""
        try:
            file_size = os.path.getsize(file_path)
            max_length = Config.ERROR_HANDLING.get('file_max_text_length', 50000)
            
            if self.should_stream(file_size):
                logger.info(f"Streaming text extraction from large file {file_path}")
                content = ""
                with open(file_path, 'r', encoding='utf-8') as f:
                    for chunk in iter(lambda: f.read(self.chunk_size), ''):
                        content += chunk
                        if len(content) > max_length:
                            logger.warning(f"Text file too large, truncating at {max_length} characters")
                            break
                return content[:max_length]
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content[:max_length]
                
        except Exception as e:
            logger.error(f"Error extracting text from TXT file {file_path}: {e}")
            return ""
    
    def _extract_docx_streaming(self, file_path: str) -> str:
        """Extract text from DOCX file with memory management"""
        try:
            file_size = os.path.getsize(file_path)
            
            # For very large DOCX files, we might want to implement additional checks
            if file_size > self.memory_threshold * 10:  # 10x the normal threshold
                logger.warning(f"DOCX file {file_path} is very large ({file_size} bytes), this may consume significant memory")
            
            doc = Document(file_path)
            
            paragraphs = []
            total_length = 0
            max_length = Config.ERROR_HANDLING.get('file_max_text_length', 50000)
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
                    total_length += len(text)
                    
                    # Prevent processing extremely large documents
                    if total_length > max_length:
                        logger.warning(f"DOCX file {file_path} too large, truncating at {max_length} characters")
                        break
            
            content = '\n'.join(paragraphs)
            return content[:max_length]
            
        except ImportError:
            logger.error("python-docx library not available for DOCX file processing")
            return ""
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
            return ""

# Global file streamer instance
file_streamer = FileStreamer()

def allowed_file(filename):
    """
    Check if file extension is allowed with comprehensive validation
    
    Args:
        filename (str): Name of the file to check
    
    Returns:
        tuple: (bool, str) - (is_valid, error_message)
    """
    if not filename:
        return False, "No filename provided"
    
    if '.' not in filename:
        return False, "File must have an extension"
    
    file_extension = filename.rsplit('.', 1)[1].lower()
    
    if file_extension not in Config.ALLOWED_EXTENSIONS:
        allowed_list = ", ".join(Config.ALLOWED_EXTENSIONS)
        return False, f"File type '{file_extension}' not allowed. Allowed types: {allowed_list}"
    
    # Check for suspicious filenames
    suspicious_patterns = ['..', '/', '\\', '<', '>', ':', '"', '|', '?', '*']
    for pattern in suspicious_patterns:
        if pattern in filename:
            return False, "Filename contains invalid characters"
    
    return True, "File type is valid"

def validate_file_upload(file):
    """
    Comprehensive file upload validation
    
    Args:
        file: Flask file upload object
    
    Returns:
        tuple: (bool, str) - (is_valid, error_message)
    """
    if not file:
        return False, "No file provided"
    
    if not file.filename:
        return False, "No file selected"
    
    # Check filename
    is_valid_name, name_error = allowed_file(file.filename)
    if not is_valid_name:
        return False, name_error
    
    # Check file size
    try:
        if not is_file_size_valid(file):
            max_size_mb = Config.MAX_CONTENT_LENGTH / (1024 * 1024)
            return False, f"File size exceeds maximum limit of {max_size_mb:.1f} MB"
    except Exception as e:
        logger.error(f"Error checking file size: {e}")
        return False, "Unable to validate file size"
    
    # Check if file is empty
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset pointer
    
    if file_size == 0:
        return False, "File is empty"
    
    if file_size < 10:  # Less than 10 bytes is likely invalid
        return False, "File appears to be invalid or corrupted"
    
    return True, "File validation passed"

def safe_get_string(data, key, default=''):
    """Safely extract and strip string values from request data, handling None values"""
    value = data.get(key, default)
    if value is None:
        return default
    return str(value).strip()

def is_file_size_valid(file):
    """
    Check if file size is within limits with error handling
    
    Args:
        file: Flask file upload object
    
    Returns:
        bool: True if file size is valid, False otherwise
    """
    try:
        current_position = file.tell()  # Save current position
        file.seek(0, 2)  # Seek to end of file
        file_size = file.tell()
        file.seek(current_position)  # Reset to original position
        return file_size <= Config.MAX_CONTENT_LENGTH
    except Exception as e:
        logger.error(f"Error checking file size: {e}")
        return False

def extract_text_from_file(file_path):
    """
    Extract text from uploaded file with comprehensive error handling
    
    Args:
        file_path (str): Path to the file
    
    Returns:
        tuple: (str or None, str) - (extracted_text, error_message)
    """
    if not file_path:
        return None, "No file path provided"
    
    if not os.path.exists(file_path):
        return None, f"File not found: {file_path}"
    
    try:
        # Check file size before processing
        file_size = os.path.getsize(file_path)
        if file_size > Config.MAX_CONTENT_LENGTH:
            max_size_mb = Config.MAX_CONTENT_LENGTH / (1024 * 1024)
            return None, f"File size ({file_size / (1024 * 1024):.1f} MB) exceeds maximum limit of {max_size_mb:.1f} MB"
        
        if file_size == 0:
            return None, "File is empty"
        
        file_extension = file_path.lower().split('.')[-1]
        
        if file_extension == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    if not text.strip():
                        return None, "Text file is empty or contains only whitespace"
                    return text, "Success"
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
                        if not text.strip():
                            return None, "Text file is empty or contains only whitespace"
                        return text, "Success"
                except Exception as e:
                    return None, f"Failed to read text file with alternative encoding: {e}"
        
        elif file_extension == 'docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                if not text_parts:
                    return None, "Word document appears to be empty or contains no readable text"
                
                text = '\n'.join(text_parts)
                return text, "Success"
                
            except Exception as e:
                return None, f"Failed to extract text from Word document: {e}"
        
        else:
            return None, f"Unsupported file type: {file_extension}"
    
    except PermissionError:
        return None, "Permission denied: Unable to access the file"
    except OSError as e:
        return None, f"Operating system error: {e}"
    except Exception as e:
        logger.error(f"Unexpected error extracting text from file {file_path}: {e}")
        return None, f"Unexpected error processing file: {e}"

def get_current_user():
    """Get current user information from session"""
    if 'user_id' not in session:
        return None
    return {
        'id': session['user_id'],
        'username': session['username'],
        'role': session['role']
    }

def login_required(f):
    """Decorator to require login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('auth.login'))
        return f(*args, **kwargs)
    return decorated_function

def role_required(role):
    """Decorator to require specific role"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'role' not in session or session['role'] != role:
                flash('Access denied. Insufficient permissions.', 'error')
                return redirect(url_for('auth.login'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def sanitize_text(text):
    """Remove or replace invalid XML characters for Word document compatibility."""
    if not isinstance(text, str):
        text = str(text)
    
    if not text:
        return ""
    
    # Remove null bytes first
    text = text.replace('\x00', '')
    
    # Replace problematic Unicode characters
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2026': '...',  # Horizontal ellipsis
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero width space
        '\u200e': '',   # Left-to-right mark
        '\u200f': '',   # Right-to-left mark
        '\ufeff': '',   # Zero width no-break space
        '—': '-',       # Em dash
        '–': '-',       # En dash
        '"': '"',       # Left double quote
        '"': '"',       # Right double quote
        ''': "'",       # Left single quote
        ''': "'",       # Right single quote
        '…': '...',     # Horizontal ellipsis
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Keep only XML-compatible characters
    result = []
    for char in text:
        cp = ord(char)
        # XML 1.0 valid characters: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
        if (cp == 0x09 or cp == 0x0A or cp == 0x0D or 
            (0x20 <= cp <= 0xD7FF) or 
            (0xE000 <= cp <= 0xFFFD) or 
            (0x10000 <= cp <= 0x10FFFF)):
            result.append(char)
        elif cp < 0x20:  # Control characters
            if cp == 0x09 or cp == 0x0A or cp == 0x0D:  # Tab, LF, CR are allowed
                result.append(char)
            # Skip other control characters
        else:
            # For other problematic characters, try to find ASCII equivalent
            if char.isspace():
                result.append(' ')
            elif char.isalnum():
                try:
                    ascii_char = char.encode('ascii', 'ignore').decode('ascii')
                    if ascii_char:
                        result.append(ascii_char)
                except:
                    pass
    
    return ''.join(result)

def extract_suggestions_from_feedback(feedback):
    """Extract suggestions from feedback containing tags <delete>, <add>, <replace>"""
    suggestions = []
    delete_pattern = re.compile(r'<delete>(.*?)</delete>')
    add_pattern = re.compile(r'<add>(.*?)</add>')
    replace_pattern = re.compile(r'<replace>(.*?)\|(.*?)</replace>')

    for match in delete_pattern.findall(feedback):
        suggestions.append({'type': 'delete', 'text': match})
    for match in add_pattern.findall(feedback):
        suggestions.append({'type': 'add', 'text': match})
    for match in replace_pattern.findall(feedback):
        suggestions.append({'type': 'replace', 'text': f'{match[0]} -> {match[1]}'})
    
    return suggestions

def create_word_document_with_suggestions(essay_text, analysis_data, accepted_suggestions):
    """Create Word document with highlighted suggestions"""
    doc = Document()
    
    # Add document properties with sanitized text
    doc.core_properties.author = sanitize_text("AI Essay Coach")
    doc.core_properties.title = sanitize_text("Essay Analysis Report")
    doc.core_properties.created = datetime.datetime.now()
    
    # Set document-wide double spacing
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    
    # Configure document styles for double spacing
    styles = doc.styles
    style = styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # Add header with logo and title
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = sanitize_text("AI Essay Coach - Analysis Report")
    header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Add score table first
    doc.add_heading(sanitize_text('Rubric Scores'), level=1)
    score_table = doc.add_table(rows=1, cols=3)  # Reduced to 3 columns (removed explanation)
    score_table.style = 'Table Grid'
    score_header_cells = score_table.rows[0].cells
    score_header_cells[0].text = sanitize_text('Category')
    score_header_cells[1].text = sanitize_text('Score')
    score_header_cells[2].text = sanitize_text('Weight')

    # Define rubric weights with proper point allocation
    rubric_weights = {
        'ideas': {'weight': 30, 'label': 'Ideas & Content'},
        'organization': {'weight': 25, 'label': 'Organization & Structure'},
        'style': {'weight': 20, 'label': 'Voice & Style'},
        'grammar': {'weight': 25, 'label': 'Grammar & Conventions'}
    }

    total_weighted_score = 0
    total_possible_weighted = 0
    
    for category, score in analysis_data['scores'].items():
        row_cells = score_table.add_row().cells
        weight_info = rubric_weights.get(category, {'weight': 25, 'label': category.title()})
        
        # Calculate actual points based on weight (score is percentage 0-100)
        actual_points = round((float(score) / 100) * weight_info['weight'], 1)
        total_weighted_score += actual_points
        total_possible_weighted += weight_info['weight']
        
        row_cells[0].text = sanitize_text(weight_info['label'])
        row_cells[1].text = sanitize_text(f"{actual_points}/{weight_info['weight']}")  # Show weighted score
        row_cells[2].text = sanitize_text(f"{weight_info['weight']}%")

    # Calculate final weighted score
    final_score = round(total_weighted_score, 1)
    
    # Add total score row
    total_row = score_table.add_row().cells
    total_row[0].text = sanitize_text('TOTAL SCORE')
    total_row[1].text = sanitize_text(f"{final_score}/100")
    total_row[2].text = sanitize_text('100%')
    
    # Make total row bold
    for cell in total_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add export timing information if available
    if 'export_timestamp' in analysis_data:
        timing_para = doc.add_paragraph()
        timing_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        timing_run = timing_para.add_run(f"Document generated on: {analysis_data['export_timestamp']}")
        timing_run.italic = True
        timing_run.font.size = Pt(9)

    # Add essay with inline suggestions and comments
    doc.add_heading(sanitize_text('Essay with AI Coaching Suggestions'), level=1)
    
    # Add instructions paragraph
    instructions = doc.add_paragraph()
    instructions.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    instructions_run = instructions.add_run(
        "Instructions: Text marked with blue strikethrough should be deleted. "
        "Text marked with red underline should be added. "
        "Explanations for each suggestion are provided in parentheses."
    )
    instructions_run.italic = True
    instructions_run.font.size = Pt(10)
    
    # Add spacing
    doc.add_paragraph()
    
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Ensure tagged_essay is sanitized
    tagged_essay = sanitize_text(str(analysis_data.get('tagged_essay', essay_text)))

    # Regex patterns for tags
    delete_pattern = re.compile(r'<delete>(.*?)</delete>')
    add_pattern = re.compile(r'<add>(.*?)</add>')
    replace_pattern = re.compile(r'<replace>(.*?)\|(.*?)</replace>')

    # Get suggestions from analysis data
    suggestions = analysis_data.get('suggestions', [])

    # Helper to add colored run with sanitized text and proper formatting
    def add_colored_run(paragraph, text, color, strike=False, underline=False):
        sanitized_text = sanitize_text(str(text))
        run = paragraph.add_run(sanitized_text)
        run.font.color.rgb = RGBColor(*color)
        run.font.strike = strike
        if underline:
            from docx.enum.text import WD_UNDERLINE
            run.font.underline = WD_UNDERLINE.SINGLE
        return run

    pos = 0
    length = len(tagged_essay)

    while pos < length:
        next_delete = delete_pattern.search(tagged_essay, pos)
        next_add = add_pattern.search(tagged_essay, pos)
        next_replace = replace_pattern.search(tagged_essay, pos)

        candidates = [m for m in [next_delete, next_add, next_replace] if m and m.start() >= pos]

        if not candidates:
            remaining_text = sanitize_text(tagged_essay[pos:])
            if remaining_text:
                paragraph.add_run(remaining_text)
            pos = length
            continue

        next_tag = min(candidates, key=lambda m: m.start())

        if next_tag.start() > pos:
            before_text = sanitize_text(tagged_essay[pos:next_tag.start()])
            if before_text:
                paragraph.add_run(before_text)

        if next_tag == next_delete:
            text = sanitize_text(str(next_tag.group(1) or ''))
            if text:
                # Deletions: blue color with strikethrough
                run = add_colored_run(paragraph, text, (0, 0, 255), strike=True)
                # Find reason for this suggestion with improved matching
                reason = ''
                for s in suggestions:
                    suggestion_type = s.get('type', '').lower()
                    suggestion_text = s.get('text', '')
                    
                    # Match deletion suggestions
                    if suggestion_type == 'delete' and text in suggestion_text:
                        reason = sanitize_text(str(s.get('reason', 'Deletion suggested for improvement')))
                        break
                    # Also check if the text appears in any suggestion
                    elif text in suggestion_text:
                        reason = sanitize_text(str(s.get('reason', 'Text modification suggested')))
                        break
                
                # Always add a reason, even if not found in suggestions
                if not reason:
                    reason = 'Remove unnecessary or incorrect text'
                
                # Add reason as inline comment (in parentheses) - smaller font, italic
                reason_run = paragraph.add_run(f' ({reason})')
                reason_run.italic = True
                reason_run.font.size = Pt(9)
                reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            pos = next_tag.end()
        elif next_tag == next_add:
            text = sanitize_text(str(next_tag.group(1) or ''))
            if text:
                # Additions: red color with underline
                run = add_colored_run(paragraph, text, (255, 0, 0), underline=True)
                # Find reason for this suggestion with improved matching
                reason = ''
                for s in suggestions:
                    suggestion_type = s.get('type', '').lower()
                    suggestion_text = s.get('text', '')
                    
                    # Match addition suggestions
                    if suggestion_type == 'add' and text in suggestion_text:
                        reason = sanitize_text(str(s.get('reason', 'Addition suggested for improvement')))
                        break
                    elif text in suggestion_text:
                        reason = sanitize_text(str(s.get('reason', 'Text addition suggested')))
                        break
                
                # Always add a reason, even if not found in suggestions
                if not reason:
                    reason = 'Add for clarity or correctness'
                
                reason_run = paragraph.add_run(f' ({reason})')
                reason_run.italic = True
                reason_run.font.size = Pt(9)
                reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            pos = next_tag.end()
        elif next_tag == next_replace:
            old_word = sanitize_text(str(next_tag.group(1) or ''))
            new_word = sanitize_text(str(next_tag.group(2) or ''))
            if old_word:
                run_old = add_colored_run(paragraph, old_word, (0, 0, 255), strike=True)
                # Find reason for replacement with improved matching
                reason = ''
                for s in suggestions:
                    suggestion_type = s.get('type', '').lower()
                    suggestion_text = s.get('text', '')
                    
                    # Match replacement suggestions
                    if suggestion_type == 'replace':
                        if (old_word in suggestion_text and new_word in suggestion_text) or \
                           f"{old_word}|{new_word}" in suggestion_text or \
                           f"{old_word} -> {new_word}" in suggestion_text:
                            reason = sanitize_text(str(s.get('reason', 'Replacement suggested for improvement')))
                            break
                    elif old_word in suggestion_text or new_word in suggestion_text:
                        reason = sanitize_text(str(s.get('reason', 'Word change suggested')))
                        break
                
                # Always add a reason, even if not found in suggestions
                if not reason:
                    reason = f'Replace "{old_word}" with "{new_word}" for better style or accuracy'
                
                reason_run = paragraph.add_run(f' ({reason})')
                reason_run.italic = True
                reason_run.font.size = Pt(9)
                reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            paragraph.add_run(' ')  # Space between old and new word
            if new_word:
                run_new = add_colored_run(paragraph, new_word, (255, 0, 0), underline=True)
            pos = next_tag.end()

    return doc

def calculate_scores_with_rubric(text_analysis, rubric_config=None):
    """Calculate scores using modular rubric engine"""
    if not rubric_config:
        # Use default weights
        rubric_config = {
            'ideas_weight': 30,
            'organization_weight': 25,
            'style_weight': 20,
            'grammar_weight': 25
        }
    
    # Extract base scores from AI analysis (0-100 scale)
    base_scores = text_analysis.get('scores', {})
    
    # Convert to weighted scores
    weighted_scores = {}
    total_score = 0
    
    for dimension in ['ideas', 'organization', 'style', 'grammar']:
        base_score = base_scores.get(dimension, 75)  # Default to 75 if missing
        weight_key = f"{dimension}_weight"
        weight = rubric_config.get(weight_key, 25)
        
        # Calculate weighted score (base_score * weight / 100)
        weighted_score = (base_score * weight) / 100
        weighted_scores[dimension] = weighted_score
        total_score += weighted_score
    
    weighted_scores['total'] = total_score
    return weighted_scores

def extract_text_from_filestorage(file_storage):
    """
    Extract text from Flask FileStorage object
    
    Args:
        file_storage: Flask FileStorage object from request.files
    
    Returns:
        str: Extracted text content
    """
    if not file_storage or not file_storage.filename:
        raise ValueError("No file provided or empty filename")
    
    filename = file_storage.filename.lower()
    
    try:
        # Read the file content into memory
        file_content = file_storage.read()
        
        # Reset the file pointer for potential future reads
        file_storage.seek(0)
        
        if not file_content:
            raise ValueError("File is empty")
        
        # Determine file type and extract text
        if filename.endswith('.txt'):
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')
                
        elif filename.endswith('.docx'):
            try:
                from docx import Document
                doc = Document(BytesIO(file_content))
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                if not text.strip():
                    raise ValueError("DOCX file contains no readable text")
                return text
            except Exception as e:
                raise ValueError(f"Failed to read DOCX file: {e}")
                
        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                text = ''
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    raise ValueError("PDF file contains no readable text")
                return text
            except Exception as e:
                raise ValueError(f"Failed to read PDF file: {e}")
                
        else:
            raise ValueError(f"Unsupported file type: {filename}")
            
    except Exception as e:
        logger.error(f"Error extracting text from FileStorage: {e}")
        raise
