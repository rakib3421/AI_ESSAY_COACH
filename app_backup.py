"""
Main Flask application for the Essay Revision System
Modularized for better maintainability
"""
from flask import Flask
import os
import logging
from config import Config
from routes import register_routes

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create Flask app
app = Flask(__name__)
app.secret_key = Config.SECRET_KEY

# Configure app settings
app.config['UPLOAD_FOLDER'] = Config.UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = Config.MAX_CONTENT_LENGTH

# Ensure upload folder exists
if not os.path.exists(Config.UPLOAD_FOLDER):
    os.makedirs(Config.UPLOAD_FOLDER)

# Register all routes
register_routes(app)

if __name__ == '__main__':
    app.run(host='10.10.12.31', port=5000, debug=True)
    connection = get_db_connection()
    if not connection:
        logger.error("Database connection failed, skipping save")
        return False
    
    try:
        with connection.cursor() as cursor:
            # Save essay analysis
            sql = """
            INSERT INTO essay_analyses (essay_text, essay_type, ideas_score, 
                                     organization_score, style_score, grammar_score, 
                                     overall_score, suggestions, word_suggestions, vocabulary_score, clarity_score, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            overall_score = sum(analysis_data['scores'].values()) // 4
            
            cursor.execute(sql, (
                essay_text,
                analysis_data['essay_type'],
                analysis_data['scores']['ideas'],
                analysis_data['scores']['organization'],
                analysis_data['scores']['style'],
                analysis_data['scores']['grammar'],
                overall_score,
                json.dumps(analysis_data.get('suggestions', [])),
                json.dumps(analysis_data.get('word_suggestions', [])),
                analysis_data.get('scores', {}).get('vocabulary', 0),
                analysis_data.get('scores', {}).get('clarity', 0),
                datetime.datetime.now()
            ))
            analysis_id = cursor.lastrowid
            analysis_data['analysis_id'] = analysis_id  # Add analysis_id to the data

            # Save examples in a separate table
            for dimension in ['ideas', 'organization', 'style', 'grammar']:
                examples = analysis_data.get('examples', {}).get(dimension, [])
                if examples and len(examples) > 0:
                    for example_text in examples[:2]:  # Limit to 2 examples
                        cursor.execute("""
                        INSERT INTO rubric_examples (analysis_id, dimension, example_text, created_at)
                        VALUES (%s, %s, %s, %s)
                        """, (analysis_id, dimension, example_text, datetime.datetime.now()))

            connection.commit()

            # Save checklist progress if provided
            checklist_progress = analysis_data.get('checklist_progress')
            student_id = analysis_data.get('student_id')
            submission_id = analysis_data.get('submission_id')
            if checklist_progress and student_id and submission_id:
                for step, completed in checklist_progress.items():
                    cursor.execute("""
                        INSERT INTO checklist_progress (student_id, submission_id, step_name, completed, updated_at)
                        VALUES (%s, %s, %s, %s, NOW())
                        ON DUPLICATE KEY UPDATE completed = VALUES(completed), updated_at = VALUES(updated_at)
                    """, (student_id, submission_id, step, bool(completed)))
                connection.commit()

            logger.info("Analysis and examples saved to database successfully")
            return analysis_id
    except Exception as e:
        logger.error(f"Database save error: {e}")
        return False
    finally:
        connection.close()

def save_submission_to_db(student_id, analysis_id, assignment_id=None):
    """Save student submission to database"""
    connection = get_db_connection()
    if not connection:
        logger.error("Database connection failed")
        return False
    
    try:
        with connection.cursor() as cursor:
            sql = """
            INSERT INTO student_submissions (student_id, analysis_id, assignment_id, submitted_at)
            VALUES (%s, %s, %s, %s)
            """
            cursor.execute(sql, (student_id, analysis_id, assignment_id, datetime.datetime.now()))
            connection.commit()
            return True
    except Exception as e:
        logger.error(f"Submission save error: {e}")
        return False
    finally:
        connection.close()

def sanitize_text(text):
    """Remove or replace invalid XML characters for Word document compatibility."""
    if not isinstance(text, str):
        text = str(text)
    
    if not text:
        return ""
    
    # Remove null bytes first
    text = text.replace('\x00', '')
    
    # Replace problematic Unicode characters
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2026': '...',  # Horizontal ellipsis
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero width space
        '\u200e': '',   # Left-to-right mark
        '\u200f': '',   # Right-to-left mark
        '\ufeff': '',   # Zero width no-break space
        '—': '-',       # Em dash
        '–': '-',       # En dash
        '"': '"',       # Left double quote
        '"': '"',       # Right double quote
        ''': "'",       # Left single quote
        ''': "'",       # Right single quote
        '…': '...',     # Horizontal ellipsis
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Keep only XML-compatible characters
    result = []
    for char in text:
        cp = ord(char)
        # XML 1.0 valid characters: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
        if (cp == 0x09 or cp == 0x0A or cp == 0x0D or 
            (0x20 <= cp <= 0xD7FF) or 
            (0xE000 <= cp <= 0xFFFD) or 
            (0x10000 <= cp <= 0x10FFFF)):
            result.append(char)
        elif cp < 0x20:  # Control characters
            if cp == 0x09 or cp == 0x0A or cp == 0x0D:  # Tab, LF, CR are allowed
                result.append(char)
            # Skip other control characters
        else:
            # For other problematic characters, try to find ASCII equivalent
            if char.isspace():
                result.append(' ')
            elif char.isalnum():
                try:
                    ascii_char = char.encode('ascii', 'ignore').decode('ascii')
                    if ascii_char:
                        result.append(ascii_char)
                except:
                    pass
    
    return ''.join(result)

def create_word_document_with_suggestions(essay_text, analysis_data, accepted_suggestions):
    doc = Document()
    
    # Add document properties with sanitized text
    doc.core_properties.author = sanitize_text("AI Essay Coach")
    doc.core_properties.title = sanitize_text("Essay Analysis Report")
    doc.core_properties.created = datetime.datetime.now()
    
    # Set document-wide double spacing
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    
    # Configure document styles for double spacing
    styles = doc.styles
    style = styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # Add header with logo and title
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = sanitize_text("AI Essay Coach - Analysis Report")
    header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Add score table first
    doc.add_heading(sanitize_text('Rubric Scores & Analysis'), level=1)
    score_table = doc.add_table(rows=1, cols=4)
    score_table.style = 'Table Grid'
    score_header_cells = score_table.rows[0].cells
    score_header_cells[0].text = sanitize_text('Category')
    score_header_cells[1].text = sanitize_text('Score')
    score_header_cells[2].text = sanitize_text('Weight')
    score_header_cells[3].text = sanitize_text('Explanation')

    # Define rubric weights
    rubric_weights = {
        'ideas': {'weight': 30, 'label': 'Ideas & Content'},
        'organization': {'weight': 25, 'label': 'Organization & Structure'},
        'style': {'weight': 20, 'label': 'Voice & Style'},
        'grammar': {'weight': 25, 'label': 'Grammar & Conventions'}
    }

    total_score = 0
    for category, score in analysis_data['scores'].items():
        row_cells = score_table.add_row().cells
        weight_info = rubric_weights.get(category, {'weight': 25, 'label': category.title()})
        row_cells[0].text = sanitize_text(weight_info['label'])
        row_cells[1].text = sanitize_text(f"{score}/{weight_info['weight']}")
        row_cells[2].text = sanitize_text(f"{weight_info['weight']}%")
        reason = analysis_data.get('score_reasons', {}).get(category, 'No reason provided.')
        row_cells[2].text = sanitize_text(str(reason))
        total_score += score

    # Add total score row
    total_row = score_table.add_row().cells
    total_row[0].text = sanitize_text('TOTAL SCORE')
    total_row[1].text = sanitize_text(f"{total_score}/100")
    total_row[2].text = sanitize_text('100%')
    total_row[3].text = sanitize_text('Overall performance based on rubric criteria')
    
    # Make total row bold
    for cell in total_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add examples section with proper formatting
    doc.add_heading(sanitize_text('Supporting Evidence for Rubric Scores'), level=1)
    for dimension in ['ideas', 'organization', 'style', 'grammar']:
        examples = analysis_data.get('examples', {}).get(dimension, [])
        if examples:
            doc.add_heading(sanitize_text(rubric_weights.get(dimension, {'label': dimension.title()})['label']), level=2)
            for i, example_text in enumerate(examples[:2], 1):  # Limit to 2 examples
                sanitized_text = sanitize_text(str(example_text))
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.add_run(f'{i}. ').bold = True
                p.add_run(sanitized_text)

    # Add essay with inline suggestions and comments
    doc.add_heading(sanitize_text('Essay with AI Coaching Suggestions'), level=1)
    
    # Add instructions paragraph
    instructions = doc.add_paragraph()
    instructions.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    instructions_run = instructions.add_run(
        "Instructions: Text marked with blue strikethrough should be deleted. "
        "Text marked with red underline should be added. "
        "Explanations for each suggestion are provided in parentheses."
    )
    instructions_run.italic = True
    instructions_run.font.size = Pt(10)
    
    # Add spacing
    doc.add_paragraph()
    
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Ensure tagged_essay is sanitized
    tagged_essay = sanitize_text(str(analysis_data.get('tagged_essay', essay_text)))

    # Regex patterns for tags
    delete_pattern = re.compile(r'<delete>(.*?)</delete>')
    add_pattern = re.compile(r'<add>(.*?)</add>')
    replace_pattern = re.compile(r'<replace>(.*?)\|(.*?)</replace>')

    # Get suggestions from analysis data
    suggestions = analysis_data.get('suggestions', [])

    # Helper to add colored run with sanitized text and proper formatting
    def add_colored_run(paragraph, text, color, strike=False, underline=False):
        sanitized_text = sanitize_text(str(text))
        run = paragraph.add_run(sanitized_text)
        run.font.color.rgb = RGBColor(*color)
        run.font.strike = strike
        if underline:
            from docx.enum.text import WD_UNDERLINE
            run.font.underline = WD_UNDERLINE.SINGLE
        return run

    pos = 0
    length = len(tagged_essay)

    while pos < length:
        next_delete = delete_pattern.search(tagged_essay, pos)
        next_add = add_pattern.search(tagged_essay, pos)
        next_replace = replace_pattern.search(tagged_essay, pos)

        candidates = [m for m in [next_delete, next_add, next_replace] if m and m.start() >= pos]

        if not candidates:
            remaining_text = sanitize_text(tagged_essay[pos:])
            if remaining_text:
                paragraph.add_run(remaining_text)
            pos = length
            continue

        next_tag = min(candidates, key=lambda m: m.start())

        if next_tag.start() > pos:
            before_text = sanitize_text(tagged_essay[pos:next_tag.start()])
            if before_text:
                paragraph.add_run(before_text)

        if next_tag == next_delete:
            text = sanitize_text(str(next_tag.group(1) or ''))
            if text:
                # Deletions: blue color with strikethrough
                run = add_colored_run(paragraph, text, (0, 0, 255), strike=True)
                # Find reason for this suggestion
                reason = ''
                for s in suggestions:
                    if s.get('type', '').lower() == 'delete' and sanitize_text(str(s.get('text', ''))) == text:
                        reason = sanitize_text(str(s.get('reason', '')))
                        break
                # Add reason as inline comment (in parentheses) - smaller font, italic
                if reason:
                    reason_run = paragraph.add_run(f' ({reason})')
                    reason_run.italic = True
                    reason_run.font.size = Pt(9)
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            pos = next_tag.end()
        elif next_tag == next_add:
            text = sanitize_text(str(next_tag.group(1) or ''))
            if text:
                # Additions: red color with underline
                run = add_colored_run(paragraph, text, (255, 0, 0), underline=True)
                reason = ''
                for s in suggestions:
                    if s.get('type', '').lower() == 'add' and sanitize_text(str(s.get('text', ''))) == text:
                        reason = sanitize_text(str(s.get('reason', '')))
                        break
                if reason:
                    reason_run = paragraph.add_run(f' ({reason})')
                    reason_run.italic = True
                    reason_run.font.size = Pt(9)
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            pos = next_tag.end()
        elif next_tag == next_replace:
            old_word = sanitize_text(str(next_tag.group(1) or ''))
            new_word = sanitize_text(str(next_tag.group(2) or ''))
            if old_word:
                run_old = add_colored_run(paragraph, old_word, (0, 0, 255), strike=True)
                reason = ''
                for s in suggestions:
                    if (s.get('type', '').lower() == 'replace' and 
                        sanitize_text(str(s.get('text', ''))) == f"{old_word} -> {new_word}"):
                        reason = sanitize_text(str(s.get('reason', '')))
                        break
                if reason:
                    reason_run = paragraph.add_run(f' ({reason})')
                    reason_run.italic = True
                    reason_run.font.size = Pt(9)
                    reason_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
            paragraph.add_run(' → ')  # Better arrow symbol
            if new_word:
                run_new = add_colored_run(paragraph, new_word, (255, 0, 0), underline=True)
            pos = next_tag.end()

    return doc

def generate_step_wise_checklist(essay_text, essay_type, analysis_data):
    """Generate step-wise checklist based on essay analysis"""
    
    base_steps = [
        {
            'name': 'Review Thesis Statement',
            'description': 'Ensure your main argument is clear and debatable',
            'order': 1,
            'required': True,
            'criteria': 'Thesis statement identified and evaluated for clarity'
        },
        {
            'name': 'Strengthen Evidence',
            'description': 'Add or improve supporting evidence for your arguments',
            'order': 2,
            'required': True,
            'criteria': 'At least two pieces of evidence per main point'
        },
        {
            'name': 'Improve Transitions',
            'description': 'Connect paragraphs and ideas more smoothly',
            'order': 3,
            'required': True,
            'criteria': 'Clear transitions between all paragraphs'
        },
        {
            'name': 'Refine Word Choice',
            'description': 'Replace weak or repetitive words with stronger alternatives',
            'order': 4,
            'required': False,
            'criteria': 'No obvious word repetition or vague language'
        },
        {
            'name': 'Polish Grammar',
            'description': 'Correct any remaining grammatical errors',
            'order': 5,
            'required': True,
            'criteria': 'No major grammatical errors remain'
        },
        {
            'name': 'Final Review',
            'description': 'Read through entire essay for flow and coherence',
            'order': 6,
            'required': True,
            'criteria': 'Essay flows smoothly from introduction to conclusion'
        }
    ]
    
    # Customize steps based on essay type
    if essay_type.lower() == 'argumentative':
        base_steps.insert(2, {
            'name': 'Address Counterarguments',
            'description': 'Acknowledge and refute opposing viewpoints',
            'order': 3,
            'required': True,
            'criteria': 'At least one counterargument addressed'
        })
    elif essay_type.lower() == 'narrative':
        base_steps[1] = {
            'name': 'Develop Characters',
            'description': 'Ensure characters are well-developed and believable',
            'order': 2,
            'required': True,
            'criteria': 'Main characters have clear motivations and development'
        }
    elif essay_type.lower() == 'literary':
        base_steps.insert(1, {
            'name': 'Check Citations',
            'description': 'Verify all quotes are properly cited and integrated',
            'order': 2,
            'required': True,
            'criteria': 'All textual evidence properly cited and explained'
        })
    
    # Unlock first step by default
    for i, step in enumerate(base_steps):
        step['unlocked'] = (i == 0)
        step['completed'] = False
    
    return base_steps

def save_step_wise_checklist(analysis_id, steps):
    """Save step-wise checklist to database"""
    connection = get_db_connection()
    if not connection:
        return False
    
    try:
        with connection.cursor() as cursor:
            for step in steps:
                cursor.execute("""
                    INSERT INTO step_wise_checklists 
                    (analysis_id, step_name, step_description, step_order, is_required, 
                     completion_criteria, is_completed, unlocked)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    analysis_id,
                    step['name'],
                    step['description'],
                    step['order'],
                    step['required'],
                    step['criteria'],
                    step['completed'],
                    step['unlocked']
                ))
            connection.commit()
            return True
    except Exception as e:
        logger.error(f"Error saving step-wise checklist: {e}")
        return False
    finally:
        connection.close()

def update_checklist_progress(analysis_id, step_name, completed=True):
    """Update checklist step completion and unlock next step if applicable"""
    connection = get_db_connection()
    if not connection:
        return False
    
    try:
        with connection.cursor() as cursor:
            # Mark current step as completed
            cursor.execute("""
                UPDATE step_wise_checklists 
                SET is_completed = %s 
                WHERE analysis_id = %s AND step_name = %s
            """, (completed, analysis_id, step_name))
            
            if completed:
                # Get current step order
                cursor.execute("""
                    SELECT step_order FROM step_wise_checklists 
                    WHERE analysis_id = %s AND step_name = %s
                """, (analysis_id, step_name))
                result = cursor.fetchone()
                
                if result:
                    current_order = result[0]
                    # Unlock next step
                    cursor.execute("""
                        UPDATE step_wise_checklists 
                        SET unlocked = TRUE 
                        WHERE analysis_id = %s AND step_order = %s
                    """, (analysis_id, current_order + 1))
            
            connection.commit()
            return True
    except Exception as e:
        logger.error(f"Error updating checklist progress: {e}")
        return False
    finally:
        connection.close()

def get_checklist_progress(analysis_id):
    """Get current checklist progress for an analysis"""
    connection = get_db_connection()
    if not connection:
        return []
    
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT step_name, step_description, step_order, is_required,
                       completion_criteria, is_completed, unlocked
                FROM step_wise_checklists
                WHERE analysis_id = %s
                ORDER BY step_order
            """, (analysis_id,))
            
            steps = []
            for row in cursor.fetchall():
                steps.append({
                    'name': row[0],
                    'description': row[1],
                    'order': row[2],
                    'required': row[3],
                    'criteria': row[4],
                    'completed': row[5],
                    'unlocked': row[6]
                })
            return steps
    except Exception as e:
        logger.error(f"Error getting checklist progress: {e}")
        return []
    finally:
        connection.close()

def create_modular_rubric_engine(teacher_id, rubric_name, weights, custom_criteria=None):
    """Create a custom rubric configuration"""
    connection = get_db_connection()
    if not connection:
        return False
    
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                INSERT INTO rubric_configurations 
                (teacher_id, name, ideas_weight, organization_weight, style_weight, 
                 grammar_weight, custom_criteria)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (
                teacher_id, rubric_name, weights['ideas'], weights['organization'],
                weights['style'], weights['grammar'], json.dumps(custom_criteria or {})
            ))
            connection.commit()
            return cursor.lastrowid
    except Exception as e:
        logger.error(f"Error creating rubric configuration: {e}")
        return False
    finally:
        connection.close()

def calculate_scores_with_rubric(text_analysis, rubric_config=None):
    """Calculate scores using modular rubric engine"""
    if not rubric_config:
        # Use default weights
        rubric_config = {
            'ideas_weight': 30,
            'organization_weight': 25,
            'style_weight': 20,
            'grammar_weight': 25
        }
    
    # Extract base scores from AI analysis (0-100 scale)
    base_scores = text_analysis.get('scores', {})
    
    # Convert to weighted scores
    weighted_scores = {}
    total_score = 0
    
    for dimension in ['ideas', 'organization', 'style', 'grammar']:
        base_score = base_scores.get(dimension, 75)  # Default to 75 if missing
        weight_key = f"{dimension}_weight"
        weight = rubric_config.get(weight_key, 25)
        
        # Calculate weighted score (base_score * weight / 100)
        weighted_score = (base_score * weight) / 100
        weighted_scores[dimension] = weighted_score
        total_score += weighted_score
    
    weighted_scores['total'] = total_score
    return weighted_scores

def get_hybrid_essay_checks(essay_text):
    """Detect and return checks for hybrid essay types"""
    try:
        content = f"""
        Analyze this essay and determine if it contains elements from multiple essay types.
        Return a JSON object with:
        - primary_type: the main essay type
        - secondary_types: list of other detected types
        - hybrid_checks: specific checks to apply for this combination
        
        Essay: {essay_text[:1000]}...
        
        Return only valid JSON.
        """
        
        if client:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": content}],
                temperature=0
            )
        else:
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": content}],
                temperature=0
            )
        
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        logger.error(f"Hybrid essay detection error: {e}")
        return {
            'primary_type': 'expository',
            'secondary_types': [],
            'hybrid_checks': []
        }

def extract_suggestions_from_feedback(feedback):
    """Extract suggestions from feedback containing tags <delete>, <add>, <replace>"""
    suggestions = []
    delete_pattern = re.compile(r'<delete>(.*?)</delete>')
    add_pattern = re.compile(r'<add>(.*?)</add>')
    replace_pattern = re.compile(r'<replace>(.*?)\|(.*?)</replace>')

    for match in delete_pattern.findall(feedback):
        suggestions.append({'type': 'delete', 'text': match})
    for match in add_pattern.findall(feedback):
        suggestions.append({'type': 'add', 'text': match})
    for match in replace_pattern.findall(feedback):
        suggestions.append({'type': 'replace', 'text': f'{match[0]} -> {match[1]}'})
    
    return suggestions

# Routes
@app.route('/')
def index():
    """Home page"""
    if 'user_id' in session:
        if session['role'] == 'student':
            return redirect(url_for('student_dashboard'))
        else:
            return redirect(url_for('teacher_dashboard'))
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, password, role FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        conn.close()
        
        if user and check_password_hash(user[2], password):
            session['user_id'] = user[0]
            session['username'] = user[1]
            session['role'] = user[3]
            
            if user[3] == 'student':
                return redirect(url_for('student_dashboard'))
            else:
                return redirect(url_for('teacher_dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    """Signup page"""
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        email = request.form['email']
        role = request.form['role']
        
        # Validate input
        if not username or not password or not email or not role:
            flash('All fields are required', 'error')
            return render_template('signup.html')
        
        if role not in ['student', 'teacher']:
            flash('Invalid role selected', 'error')
            return render_template('signup.html')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Check if username exists
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        if cursor.fetchone():
            flash('Username already exists', 'error')
            conn.close()
            return render_template('signup.html')
        
        # Check if email exists
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        if cursor.fetchone():
            flash('Email already exists', 'error')
            conn.close()
            return render_template('signup.html')
        
        # Create user
        hashed_password = generate_password_hash(password)
        cursor.execute(
            "INSERT INTO users (username, password, email, role, created_at) VALUES (%s, %s, %s, %s, %s)",
            (username, hashed_password, email, role, datetime.datetime.now())
        )
        conn.commit()
        
        # Get user ID
        user_id = cursor.lastrowid
        conn.close()
        
        # Set session
        session['user_id'] = user_id
        session['username'] = username
        session['role'] = role
        
        if role == 'student':
            return redirect(url_for('student_dashboard'))
        else:
            return redirect(url_for('teacher_dashboard'))
    
    return render_template('signup.html')

@app.route('/logout')
def logout():
    """Logout user"""
    session.clear()
    flash('Logged out successfully', 'success')
    return redirect(url_for('index'))

@app.route('/student/dashboard')
@login_required
@role_required('student')
def student_dashboard():
    """Student dashboard"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get recent submissions
    cursor.execute("""
        SELECT id, title, essay_type, status, created_at, total_score, teacher_feedback 
        FROM essays 
        WHERE user_id = %s 
        ORDER BY created_at DESC 
        LIMIT 5
    """, (session['user_id'],))
    recent_submissions = cursor.fetchall()
    
    # Get all essays for recent activity and stats
    cursor.execute("""
        SELECT id, title, essay_type, status, created_at, total_score, teacher_feedback 
        FROM essays 
        WHERE user_id = %s 
        ORDER BY created_at DESC
    """, (session['user_id'],))
    all_essays = cursor.fetchall()
    
    # Calculate statistics
    total_essays = len(all_essays)
    scored_essays = [essay for essay in all_essays if essay[5] is not None]
    average_score = sum(essay[5] for essay in scored_essays) / len(scored_essays) if scored_essays else 0
    
    # Get progress data
    cursor.execute("""
        SELECT AVG(ideas_score) as avg_ideas, AVG(organization_score) as avg_org, 
               AVG(style_score) as avg_style, AVG(grammar_score) as avg_grammar
        FROM essays 
        WHERE user_id = %s AND total_score IS NOT NULL
    """, (session['user_id'],))
    progress = cursor.fetchone()
    
    # Get assignments with feedback info
    cursor.execute("""
        SELECT a.id, a.title, a.description, a.essay_type, a.due_date,
               CASE WHEN s.id IS NOT NULL THEN 'Submitted' ELSE 'Pending' END as status,
               s.essay_id,
               e.teacher_feedback
        FROM assignments a
        LEFT JOIN assignment_submissions s ON a.id = s.assignment_id AND s.student_id = %s
        LEFT JOIN essays e ON s.essay_id = e.id
        WHERE a.teacher_id IN (
            SELECT teacher_id FROM student_teacher_assignments WHERE student_id = %s
        )
        ORDER BY a.due_date
    """, (session['user_id'], session['user_id']))
    assignments = cursor.fetchall()
    
    pending_assignments = len([a for a in assignments if a[5] == 'Pending'])
    
    # Get pending assignment requests count
    cursor.execute("""
        SELECT COUNT(*) FROM assignment_requests 
        WHERE student_id = %s AND status = 'pending'
    """, (session['user_id'],))
    pending_requests_count = cursor.fetchone()[0]
    
    conn.close()
    
    return render_template('student/dashboard.html', 
                         submissions=recent_submissions, 
                         recent_essays=all_essays[:5],  # For recent activity section
                         total_essays=total_essays,
                         average_score=average_score,
                         pending_assignments=pending_assignments,
                         days_streak=0,  # Placeholder for now
                         progress=progress, 
                         assignments=assignments,
                         pending_requests_count=pending_requests_count)

@app.route('/student/upload', methods=['GET', 'POST'])
@login_required
@role_required('student')
def upload_essay():
    """Upload essay for analysis"""
    assignment_id = request.args.get('assignment_id', type=int)
    assignment_essay_type = None
    
    # If this is for an assignment, get the assignment's essay type
    if assignment_id:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT essay_type FROM assignments WHERE id = %s", (assignment_id,))
        assignment_data = cursor.fetchone()
        conn.close()
        if assignment_data:
            assignment_essay_type = assignment_data[0]
    
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        title = request.form['title']
        essay_type = request.form.get('essay_type', 'auto')  # Default to 'auto' if not selected
        coaching_level = request.form.get('coaching_level', 'medium')
        suggestion_level = request.form.get('suggestion_level', 'medium')
        
        # Get assignment_id from form if it exists (overrides URL parameter)
        form_assignment_id = request.form.get('assignment_id', type=int)
        if form_assignment_id:
            assignment_id = form_assignment_id
        
        # If this is for an assignment and the assignment has a specific type (not 'auto'), use it
        if assignment_id and assignment_essay_type and assignment_essay_type != 'auto':
            essay_type = assignment_essay_type
        
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            # Check file size
            if not is_file_size_valid(file):
                flash('File too large. Maximum size is 16MB.', 'error')
                return redirect(request.url)
                
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Extract text from file
            try:
                text = extract_text_from_file(file_path)
                if not text:
                    flash('Could not extract text from file. Please check the file format.', 'error')
                    os.remove(file_path)
                    return redirect(request.url)
                
                # Store analysis data and redirect to new view
                analysis_data = {
                    'essay': text,
                    'title': title,
                    'essay_type': essay_type,
                    'coaching_level': coaching_level,
                    'suggestion_aggressiveness': suggestion_level,
                    'assignment_id': assignment_id  # Include assignment_id if it exists
                }
                
                # Store in session storage for the view page
                session['temp_analysis_data'] = analysis_data
                
                # Clean up uploaded file
                os.remove(file_path)
                
                flash('Essay uploaded successfully! Analyzing with AI...', 'success')
                return redirect(url_for('analyze_view'))
                
            except Exception as e:
                flash(f'Error processing file: {str(e)}', 'error')
                if os.path.exists(file_path):
                    os.remove(file_path)
        else:
            flash('Invalid file type. Please upload a .docx or .txt file.', 'error')
    
    # Get assignment details to pass to template
    assignment = None
    if assignment_id:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id, title, essay_type, description FROM assignments WHERE id = %s", (assignment_id,))
        assignment = cursor.fetchone()
        conn.close()
    
    return render_template('student/upload_new.html', assignment=assignment)

@app.route('/student/essay/<int:essay_id>')
@login_required
@role_required('student')
def view_essay(essay_id):
    """View essay with AI suggestions"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, title, content, essay_type, feedback, ideas_score, 
               organization_score, style_score, grammar_score, total_score, status
        FROM essays 
        WHERE id = %s AND user_id = %s
    """, (essay_id, session['user_id']))
    essay = cursor.fetchone()
    conn.close()
    
    if not essay:
        flash('Essay not found', 'error')
        return redirect(url_for('student_dashboard'))
    
    return render_template('student/view_essay.html', essay=essay)

@app.route('/teacher/dashboard')
@login_required
@role_required('teacher')
def teacher_dashboard():
    """Teacher dashboard"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get student submissions from assigned students only
    cursor.execute("""
        SELECT e.id, e.title, e.essay_type, e.total_score, e.created_at, u.username
        FROM essays e
        JOIN users u ON e.user_id = u.id
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE u.role = 'student' AND sta.teacher_id = %s
        ORDER BY e.created_at DESC
        LIMIT 10
    """, (session['user_id'],))
    submissions = cursor.fetchall()
    
    # Get assignments created by this teacher
    cursor.execute("""
        SELECT id, title, essay_type, due_date, created_at
        FROM assignments
        WHERE teacher_id = %s
        ORDER BY created_at DESC
        LIMIT 5
    """, (session['user_id'],))
    assignments = cursor.fetchall()
    
    # Get assigned students count
    cursor.execute("""
        SELECT COUNT(*) FROM student_teacher_assignments WHERE teacher_id = %s
    """, (session['user_id'],))
    assigned_students_count = cursor.fetchone()[0]
    
    # Get count of essays needing review (no teacher feedback)
    cursor.execute("""
        SELECT COUNT(*) FROM essays e
        JOIN users u ON e.user_id = u.id
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE u.role = 'student' AND sta.teacher_id = %s AND e.teacher_feedback IS NULL
    """, (session['user_id'],))
    pending_reviews = cursor.fetchone()[0]
    
    conn.close()
    
    return render_template('teacher/dashboard.html', 
                         submissions=submissions, 
                         assignments=assignments,
                         assigned_students_count=assigned_students_count,
                         active_students=assigned_students_count,  # For template compatibility
                         pending_reviews=pending_reviews)

@app.route('/teacher/submissions')
@login_required
@role_required('teacher')
def view_submissions():
    """View all student submissions from assigned students"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get all submissions from assigned students with pagination support
    cursor.execute("""
        SELECT e.id, e.title, e.essay_type, e.total_score, e.created_at, u.username, e.status
        FROM essays e
        JOIN users u ON e.user_id = u.id
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE u.role = 'student' AND sta.teacher_id = %s
        ORDER BY e.created_at DESC
        LIMIT 50
    """, (session['user_id'],))
    submissions = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/submissions.html', submissions=submissions)

@app.route('/teacher/create_assignment', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def create_assignment():
    """Create new assignment"""
    if request.method == 'POST':
        title = request.form['title']
        description = request.form['description']
        essay_type = request.form.get('essay_type', 'auto')  # Default to 'auto' if not selected
        due_date = request.form['due_date']
        guidelines = request.form['guidelines']
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO assignments (teacher_id, title, description, essay_type, due_date, guidelines, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (session['user_id'], title, description, essay_type, due_date, guidelines, datetime.datetime.now()))
        conn.commit()
        conn.close()
        
        flash('Assignment created successfully!', 'success')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/create_assignment.html')

@app.route('/student/essays')
@login_required
@role_required('student')
def student_essays_list():
    """View all essays submitted by student"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get all essays by the student
        cursor.execute("""
            SELECT id, title, essay_type, total_score, created_at, status, teacher_feedback,
                   ideas_score, organization_score, style_score, grammar_score
            FROM essays 
            WHERE user_id = %s 
            ORDER BY created_at DESC
        """, (session['user_id'],))
        essays = cursor.fetchall()
        
        # Get assignment submissions for context
        cursor.execute("""
            SELECT e.id, a.title as assignment_title, a.id as assignment_id
            FROM essays e
            JOIN assignment_submissions asub ON e.id = asub.essay_id
            JOIN assignments a ON asub.assignment_id = a.id
            WHERE e.user_id = %s
        """, (session['user_id'],))
        assignment_context = {row[0]: {'title': row[1], 'id': row[2]} for row in cursor.fetchall()}
        
        conn.close()
        
        return render_template('student/essays.html', essays=essays, assignment_context=assignment_context)
        
    except Exception as e:
        logger.error(f"Error in student_essays_list: {e}")
        flash('Error loading essays', 'error')
        return redirect(url_for('student_dashboard'))

@app.route('/student/assignments')
@login_required
@role_required('student')
def view_assignments():
    """View all assignments for student"""
    try:
        conn = get_db_connection()
        if not conn:
            logger.error("Database connection failed")
            flash('Database connection error. Please try again later.', 'error')
            return redirect(url_for('student_dashboard'))
        
        cursor = conn.cursor()
        
        # First check if student has any teacher assignments
        cursor.execute("""
            SELECT COUNT(*) FROM student_teacher_assignments WHERE student_id = %s
        """, (session['user_id'],))
        teacher_count = cursor.fetchone()[0]
        
        logger.info(f"Student {session['user_id']} has {teacher_count} teacher assignments")
        
        if teacher_count == 0:
            # No teachers assigned, show all assignments for now (or create demo data)
            cursor.execute("""
                SELECT a.id, a.title, a.description, a.essay_type, a.due_date, a.guidelines,
                       CASE WHEN s.id IS NOT NULL THEN 'Submitted' ELSE 'Pending' END as status,
                       s.essay_id,
                       e.teacher_feedback
                FROM assignments a
                LEFT JOIN assignment_submissions s ON a.id = s.assignment_id AND s.student_id = %s
                LEFT JOIN essays e ON s.essay_id = e.id
                ORDER BY a.due_date
            """, (session['user_id'],))
        else:
            # Normal query for students with assigned teachers
            cursor.execute("""
                SELECT a.id, a.title, a.description, a.essay_type, a.due_date, a.guidelines,
                       CASE WHEN s.id IS NOT NULL THEN 'Submitted' ELSE 'Pending' END as status,
                       s.essay_id,
                       e.teacher_feedback
                FROM assignments a
                LEFT JOIN assignment_submissions s ON a.id = s.assignment_id AND s.student_id = %s
                LEFT JOIN essays e ON s.essay_id = e.id
                WHERE a.teacher_id IN (
                    SELECT teacher_id FROM student_teacher_assignments WHERE student_id = %s
                )
                ORDER BY a.due_date
            """, (session['user_id'], session['user_id']))
        
        assignments = cursor.fetchall()
        logger.info(f"Found {len(assignments)} assignments for student {session['user_id']}")
        
        conn.close()
        
        return render_template('student/assignments.html', assignments=assignments)
        
    except Exception as e:
        logger.error(f"Error in view_assignments: {e}")
        flash('Error loading assignments. Please try again.', 'error')
        return redirect(url_for('student_dashboard'))

@app.route('/student/analyze-view')
@login_required
@role_required('student')
def analyze_view():
    """Direct analysis view for text input"""
    # Clear temporary session data after using it
    temp_data = session.pop('temp_analysis_data', None)
    return render_template('student/view_essay_new.html', essay=None, temp_data=temp_data)

@app.route('/student/submit_assignment/<int:assignment_id>')
@login_required
@role_required('student')
def submit_assignment(assignment_id):
    """Submit assignment"""
    # This would redirect to upload with assignment context
    return redirect(url_for('upload_essay', assignment_id=assignment_id))

@app.route('/api/essay/<int:essay_id>/suggestions')
@login_required
def get_essay_suggestions(essay_id):
    """Get AI suggestions for an essay"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Verify user owns essay or is teacher
    if session['role'] == 'student':
        cursor.execute("SELECT id FROM essays WHERE id = %s AND user_id = %s", (essay_id, session['user_id']))
    else:
        cursor.execute("""
            SELECT e.id FROM essays e
            JOIN users u ON e.user_id = u.id
            JOIN student_teacher_assignments sta ON u.id = sta.student_id
            WHERE e.id = %s AND sta.teacher_id = %s
        """, (essay_id, session['user_id']))
    
    if not cursor.fetchone():
        conn.close()
        return jsonify({'error': 'Essay not found'}), 404
    
    # Get suggestions
    cursor.execute("""
        SELECT id, original_text, suggested_text, suggestion_type, explanation, 
               position_start, position_end, is_accepted
        FROM essay_suggestions 
        WHERE essay_id = %s 
        ORDER BY position_start
    """, (essay_id,))
    suggestions = cursor.fetchall()
    conn.close()
    
    return jsonify({
        'suggestions': [{
            'id': s[0],
            'original_text': s[1],
            'suggested_text': s[2],
            'type': s[3],
            'explanation': s[4],
            'position_start': s[5],
            'position_end': s[6],
            'is_accepted': s[7]
        } for s in suggestions]
    })

@app.route('/api/suggestion/<int:suggestion_id>/accept', methods=['POST'])
@login_required
def accept_suggestion(suggestion_id):
    """Accept a suggestion"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("UPDATE essay_suggestions SET is_accepted = TRUE WHERE id = %s", (suggestion_id,))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/suggestion/<int:suggestion_id>/reject', methods=['POST'])
@login_required
def reject_suggestion(suggestion_id):
    """Reject a suggestion"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM essay_suggestions WHERE id = %s", (suggestion_id,))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/teacher/feedback/<int:essay_id>', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def provide_feedback(essay_id):
    """Provide teacher feedback on essay"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        feedback = request.form['feedback']
        
        # Get score overrides if provided
        ideas_score = request.form.get('ideas_score')
        organization_score = request.form.get('organization_score')
        style_score = request.form.get('style_score')
        grammar_score = request.form.get('grammar_score')
        
        # Convert empty strings to None for database
        ideas_score = float(ideas_score) if ideas_score else None
        organization_score = float(organization_score) if organization_score else None
        style_score = float(style_score) if style_score else None
        grammar_score = float(grammar_score) if grammar_score else None
        
        # Calculate total score if individual scores are provided
        total_score = None
        if all(score is not None for score in [ideas_score, organization_score, style_score, grammar_score]):
            total_score = ideas_score + organization_score + style_score + grammar_score
        
        # Update essay with feedback and scores
        update_query = """
            UPDATE essays SET 
                teacher_feedback = %s, 
                status = 'reviewed'
        """
        update_params = [feedback]
        
        if ideas_score is not None:
            update_query += ", ideas_score = %s"
            update_params.append(ideas_score)
        if organization_score is not None:
            update_query += ", organization_score = %s"
            update_params.append(organization_score)
        if style_score is not None:
            update_query += ", style_score = %s"
            update_params.append(style_score)
        if grammar_score is not None:
            update_query += ", grammar_score = %s"
            update_params.append(grammar_score)
        if total_score is not None:
            update_query += ", total_score = %s"
            update_params.append(total_score)
            
        update_query += " WHERE id = %s"
        update_params.append(essay_id)
        
        cursor.execute(update_query, update_params)
        conn.commit()
        flash('Feedback provided successfully!', 'success')
        return redirect(url_for('teacher_dashboard'))
    
    # Get essay details - Updated query to work without student-teacher assignments
    cursor.execute("""
        SELECT e.id, e.title, e.content, e.essay_type, e.feedback, e.teacher_feedback,
               e.ideas_score, e.organization_score, e.style_score, e.grammar_score, 
               e.total_score, u.username
        FROM essays e
        JOIN users u ON e.user_id = u.id
        WHERE e.id = %s AND u.role = 'student'
    """, (essay_id,))
    essay = cursor.fetchone()
    conn.close()
    
    if not essay:
        flash('Essay not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/feedback.html', essay=essay)

@app.route('/teacher/assignment/<int:assignment_id>')
@login_required
@role_required('teacher')
def view_assignment(assignment_id):
    """View assignment details"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get assignment details
    cursor.execute("""
        SELECT id, title, description, essay_type, due_date, guidelines, created_at
        FROM assignments
        WHERE id = %s AND teacher_id = %s
    """, (assignment_id, session['user_id']))
    assignment = cursor.fetchone()
    
    if not assignment:
        flash('Assignment not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    # Get submission statistics
    cursor.execute("""
        SELECT COUNT(*) as total_submissions,
               AVG(e.total_score) as avg_score
        FROM assignment_submissions asub
        JOIN essays e ON asub.essay_id = e.id
        WHERE asub.assignment_id = %s
    """, (assignment_id,))
    stats = cursor.fetchone()
    
    conn.close()
    
    return render_template('teacher/assignment_view.html', assignment=assignment, stats=stats)

@app.route('/teacher/assignment/<int:assignment_id>/submissions')
@login_required
@role_required('teacher')
def view_assignment_submissions(assignment_id):
    """View all submissions for an assignment"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get assignment details
    cursor.execute("""
        SELECT title, description, essay_type, due_date
        FROM assignments
        WHERE id = %s AND teacher_id = %s
    """, (assignment_id, session['user_id']))
    assignment = cursor.fetchone()
    
    if not assignment:
        flash('Assignment not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    # Get all submissions for this assignment
    cursor.execute("""
        SELECT e.id, e.title, e.total_score, e.created_at, u.username, 
               e.teacher_feedback, e.status
        FROM assignment_submissions asub
        JOIN essays e ON asub.essay_id = e.id
        JOIN users u ON e.user_id = u.id
        WHERE asub.assignment_id = %s
        ORDER BY e.created_at DESC
    """, (assignment_id,))
    submissions = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/assignment_submissions.html', 
                         assignment=assignment, submissions=submissions, assignment_id=assignment_id)

@app.route('/teacher/assignment/<int:assignment_id>/edit', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def edit_assignment(assignment_id):
    """Edit assignment"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        title = request.form['title']
        description = request.form['description']
        essay_type = request.form.get('essay_type', 'auto')  # Default to 'auto' if not selected
        due_date = request.form['due_date']
        guidelines = request.form['guidelines']
        
        cursor.execute("""
            UPDATE assignments 
            SET title = %s, description = %s, essay_type = %s, due_date = %s, guidelines = %s
            WHERE id = %s AND teacher_id = %s
        """, (title, description, essay_type, due_date, guidelines, assignment_id, session['user_id']))
        conn.commit()
        conn.close()
        
        flash('Assignment updated successfully!', 'success')
        return redirect(url_for('teacher_dashboard'))
    
    # Get assignment details
    cursor.execute("""
        SELECT id, title, description, essay_type, due_date, guidelines
        FROM assignments
        WHERE id = %s AND teacher_id = %s
    """, (assignment_id, session['user_id']))
    assignment = cursor.fetchone()
    conn.close()
    
    if not assignment:
        flash('Assignment not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/edit_assignment.html', assignment=assignment)

@app.route('/teacher/assignment/<int:assignment_id>/feedback/<int:essay_id>', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def provide_assignment_feedback(assignment_id, essay_id):
    """Provide teacher feedback on assignment submission"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        feedback = request.form['feedback']
        
        # Get score overrides if provided
        ideas_score = request.form.get('ideas_score')
        organization_score = request.form.get('organization_score')
        style_score = request.form.get('style_score')
        grammar_score = request.form.get('grammar_score')
        
        # Convert empty strings to None for database
        ideas_score = float(ideas_score) if ideas_score else None
        organization_score = float(organization_score) if organization_score else None
        style_score = float(style_score) if style_score else None
        grammar_score = float(grammar_score) if grammar_score else None
        
        # Calculate total score if individual scores are provided
        total_score = None
        if all(score is not None for score in [ideas_score, organization_score, style_score, grammar_score]):
            total_score = ideas_score + organization_score + style_score + grammar_score
        
        # Update essay with feedback and scores
        update_query = """
            UPDATE essays SET 
                teacher_feedback = %s, 
                status = 'reviewed'
        """
        update_params = [feedback]
        
        if ideas_score is not None:
            update_query += ", ideas_score = %s"
            update_params.append(ideas_score)
        if organization_score is not None:
            update_query += ", organization_score = %s"
            update_params.append(organization_score)
        if style_score is not None:
            update_query += ", style_score = %s"
            update_params.append(style_score)
        if grammar_score is not None:
            update_query += ", grammar_score = %s"
            update_params.append(grammar_score)
        if total_score is not None:
            update_query += ", total_score = %s"
            update_params.append(total_score)
            
        update_query += " WHERE id = %s"
        update_params.append(essay_id)
        
        cursor.execute(update_query, update_params)
        conn.commit()
        flash('Feedback provided successfully!', 'success')
        return redirect(url_for('view_assignment_submissions', assignment_id=assignment_id))
    
    # Get essay details
    cursor.execute("""
        SELECT e.id, e.title, e.content, e.essay_type, e.feedback, e.teacher_feedback,
               e.ideas_score, e.organization_score, e.style_score, e.grammar_score, 
               e.total_score, u.username
        FROM essays e
        JOIN users u ON e.user_id = u.id
        WHERE e.id = %s AND u.role = 'student'
    """, (essay_id,))
    essay = cursor.fetchone()
    
    # Get assignment details for context
    cursor.execute("""
        SELECT title, essay_type, due_date 
        FROM assignments 
        WHERE id = %s AND teacher_id = %s
    """, (assignment_id, session['user_id']))
    assignment = cursor.fetchone()
    
    conn.close()
    
    if not essay or not assignment:
        flash('Essay or assignment not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/assignment_feedback.html', essay=essay, assignment=assignment, assignment_id=assignment_id)

@app.route('/teacher/students')
@login_required
@role_required('teacher')
def student_management():
    """Student management page for teachers"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get all students
    cursor.execute("""
        SELECT u.id, u.username, u.email, u.created_at,
               CASE WHEN sta.id IS NOT NULL THEN 'assigned' ELSE 'not_assigned' END as status
        FROM users u
        LEFT JOIN student_teacher_assignments sta ON u.id = sta.student_id AND sta.teacher_id = %s
        WHERE u.role = 'student'
        ORDER BY u.username
    """, (session['user_id'],))
    all_students = cursor.fetchall()
    
    # Get assigned students with statistics
    cursor.execute("""
        SELECT u.id, u.username, u.email, sta.assigned_at,
               COUNT(DISTINCT e.id) as essay_count,
               COUNT(DISTINCT asub.id) as assignment_submissions,
               AVG(e.total_score) as avg_score
        FROM users u
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        LEFT JOIN essays e ON u.id = e.user_id
        LEFT JOIN assignment_submissions asub ON u.id = asub.student_id 
            AND asub.assignment_id IN (SELECT id FROM assignments WHERE teacher_id = %s)
        WHERE sta.teacher_id = %s AND u.role = 'student'
        GROUP BY u.id, u.username, u.email, sta.assigned_at
        ORDER BY u.username
    """, (session['user_id'], session['user_id']))
    assigned_students = cursor.fetchall()
    
    # Get pending requests
    cursor.execute("""
        SELECT ar.id, u.username, u.email, ar.created_at, ar.status
        FROM assignment_requests ar
        JOIN users u ON ar.student_id = u.id
        WHERE ar.teacher_id = %s AND ar.status = 'pending'
        ORDER BY ar.created_at DESC
    """, (session['user_id'],))
    pending_requests = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/students.html', 
                         all_students=all_students, 
                         assigned_students=assigned_students,
                         pending_requests=pending_requests)

@app.route('/teacher/all_essays')
@login_required
@role_required('teacher')
def teacher_all_essays():
    """View all essays from assigned students with feedback options"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get all essays from assigned students
    cursor.execute("""
        SELECT e.id, e.title, e.essay_type, e.total_score, e.created_at, e.status, 
               e.teacher_feedback, u.username, e.ideas_score, e.organization_score, 
               e.style_score, e.grammar_score
        FROM essays e
        JOIN users u ON e.user_id = u.id
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE sta.teacher_id = %s AND u.role = 'student'
        ORDER BY e.created_at DESC
    """, (session['user_id'],))
    essays = cursor.fetchall()
    
    # Get assignment context for each essay
    cursor.execute("""
        SELECT e.id, a.title as assignment_title, a.id as assignment_id
        FROM essays e
        JOIN assignment_submissions asub ON e.id = asub.essay_id
        JOIN assignments a ON asub.assignment_id = a.id
        WHERE a.teacher_id = %s
    """, (session['user_id'],))
    assignment_context = {row[0]: {'title': row[1], 'id': row[2]} for row in cursor.fetchall()}
    
    conn.close()
    
    return render_template('teacher/all_essays.html', essays=essays, assignment_context=assignment_context)

@app.route('/teacher/students/<int:student_id>/request', methods=['POST'])
@login_required
@role_required('teacher')
def send_assignment_request(student_id):
    """Send assignment request to student"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Check if student exists
    cursor.execute("SELECT id FROM users WHERE id = %s AND role = 'student'", (student_id,))
    if not cursor.fetchone():
        flash('Student not found', 'error')
        return redirect(url_for('student_management'))
    
    # Check if already assigned
    cursor.execute("""
        SELECT id FROM student_teacher_assignments 
        WHERE student_id = %s AND teacher_id = %s
    """, (student_id, session['user_id']))
    if cursor.fetchone():
        flash('Student is already assigned to you', 'warning')
        return redirect(url_for('student_management'))
    
    # Check if request already exists
    cursor.execute("""
        SELECT id FROM assignment_requests 
        WHERE student_id = %s AND teacher_id = %s AND status = 'pending'
    """, (student_id, session['user_id']))
    if cursor.fetchone():
        flash('Assignment request already sent', 'warning')
        return redirect(url_for('student_management'))
    
    # Create assignment request
    cursor.execute("""
        INSERT INTO assignment_requests (teacher_id, student_id, status, created_at)
        VALUES (%s, %s, 'pending', %s)
    """, (session['user_id'], student_id, datetime.datetime.now()))
    conn.commit()
    conn.close()
    
    flash('Assignment request sent successfully!', 'success')
    return redirect(url_for('student_management'))

@app.route('/teacher/students/<int:student_id>/essays')
@login_required
@role_required('teacher')
def view_student_essays(student_id):
    """View essays from a specific assigned student"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Verify student is assigned to this teacher
    cursor.execute("""
        SELECT u.username FROM users u
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE u.id = %s AND sta.teacher_id = %s AND u.role = 'student'
    """, (student_id, session['user_id']))
    student = cursor.fetchone()
    
    if not student:
        flash('Student not found or not assigned to you', 'error')
        return redirect(url_for('student_management'))
    
    # Get student's essays
    cursor.execute("""
        SELECT id, title, essay_type, total_score, created_at, status, teacher_feedback
        FROM essays
        WHERE user_id = %s
        ORDER BY created_at DESC
    """, (student_id,))
    essays = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/student_essays.html', 
                         student_name=student[0], 
                         student_id=student_id,
                         essays=essays)

@app.route('/teacher/students/<int:student_id>/assignments')
@login_required
@role_required('teacher')
def view_student_assignments(student_id):
    """View assignment submissions from a specific assigned student"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Verify student is assigned to this teacher
    cursor.execute("""
        SELECT u.username FROM users u
        JOIN student_teacher_assignments sta ON u.id = sta.student_id
        WHERE u.id = %s AND sta.teacher_id = %s AND u.role = 'student'
    """, (student_id, session['user_id']))
    student = cursor.fetchone()
    
    if not student:
        flash('Student not found or not assigned to you', 'error')
        return redirect(url_for('student_management'))
    
    # Get student's assignment submissions for this teacher's assignments
    cursor.execute("""
        SELECT a.id, a.title, a.essay_type, a.due_date,
               CASE WHEN asub.id IS NOT NULL THEN 'Submitted' ELSE 'Pending' END as status,
               e.id as essay_id, e.total_score, e.teacher_feedback, asub.submitted_at
        FROM assignments a
        LEFT JOIN assignment_submissions asub ON a.id = asub.assignment_id AND asub.student_id = %s
        LEFT JOIN essays e ON asub.essay_id = e.id
        WHERE a.teacher_id = %s
        ORDER BY a.due_date DESC
    """, (student_id, session['user_id']))
    assignments = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/student_assignments.html', 
                         student_name=student[0], 
                         student_id=student_id,
                         assignments=assignments)

@app.route('/student/requests')
@login_required
@role_required('student')
def view_assignment_requests():
    """View assignment requests from teachers"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get pending requests
    cursor.execute("""
        SELECT ar.id, u.username, u.email, ar.created_at
        FROM assignment_requests ar
        JOIN users u ON ar.teacher_id = u.id
        WHERE ar.student_id = %s AND ar.status = 'pending'
        ORDER BY ar.created_at DESC
    """, (session['user_id'],))
    requests = cursor.fetchall()
    
    conn.close()
    
    return render_template('student/assignment_requests.html', requests=requests)

@app.route('/student/requests/<int:request_id>/accept', methods=['POST'])
@login_required
@role_required('student')
def accept_assignment_request(request_id):
    """Accept assignment request from teacher"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get request details
    cursor.execute("""
        SELECT teacher_id, student_id FROM assignment_requests
        WHERE id = %s AND student_id = %s AND status = 'pending'
    """, (request_id, session['user_id']))
    request_data = cursor.fetchone()
    
    if not request_data:
        flash('Request not found', 'error')
        return redirect(url_for('view_assignment_requests'))
    
    teacher_id, student_id = request_data
    
    # Check if already assigned
    cursor.execute("""
        SELECT id FROM student_teacher_assignments
        WHERE student_id = %s AND teacher_id = %s
    """, (student_id, teacher_id))
    if cursor.fetchone():
        # Update request status
        cursor.execute("""
            UPDATE assignment_requests SET status = 'accepted'
            WHERE id = %s
        """, (request_id,))
        conn.commit()
        flash('You are already assigned to this teacher', 'info')
    else:
        # Create assignment
        cursor.execute("""
            INSERT INTO student_teacher_assignments (student_id, teacher_id, assigned_at)
            VALUES (%s, %s, %s)
        """, (student_id, teacher_id, datetime.datetime.now()))
        
        # Update request status
        cursor.execute("""
            UPDATE assignment_requests SET status = 'accepted'
            WHERE id = %s
        """, (request_id,))
        
        conn.commit()
        flash('Assignment request accepted! You are now assigned to this teacher.', 'success')
    
    conn.close()
    return redirect(url_for('view_assignment_requests'))

@app.route('/student/requests/<int:request_id>/reject', methods=['POST'])
@login_required
@role_required('student')
def reject_assignment_request(request_id):
    """Reject assignment request from teacher"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Update request status
    cursor.execute("""
        UPDATE assignment_requests SET status = 'rejected'
        WHERE id = %s AND student_id = %s AND status = 'pending'
    """, (request_id, session['user_id']))
    
    if cursor.rowcount > 0:
        conn.commit()
        flash('Assignment request rejected', 'info')
    else:
        flash('Request not found', 'error')
    
    conn.close()
    return redirect(url_for('view_assignment_requests'))

@app.route('/student/progress')
@login_required
@role_required('student')
def student_progress():
    """View detailed student progress"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get all essays with scores
    cursor.execute("""
        SELECT id, title, essay_type, ideas_score, organization_score, 
               style_score, grammar_score, total_score, created_at, teacher_feedback
        FROM essays 
        WHERE user_id = %s AND total_score IS NOT NULL
        ORDER BY created_at
    """, (session['user_id'],))
    essays = cursor.fetchall()
    
    # Get progress over time
    cursor.execute("""
        SELECT DATE(created_at) as date, AVG(total_score) as avg_score
        FROM essays 
        WHERE user_id = %s AND total_score IS NOT NULL
        GROUP BY DATE(created_at)
        ORDER BY date
    """, (session['user_id'],))
    progress_data = cursor.fetchall()
    
    conn.close()
    
    return render_template('student/progress.html', essays=essays, progress_data=progress_data)

@app.route('/student/feedback/<int:essay_id>')
@login_required
@role_required('student')
def view_feedback(essay_id):
    """View teacher feedback for an essay"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get essay details including teacher feedback
    cursor.execute("""
        SELECT e.id, e.title, e.content, e.essay_type, e.feedback, e.teacher_feedback,
               e.ideas_score, e.organization_score, e.style_score, e.grammar_score, 
               e.total_score, e.created_at, u.username as teacher_name
        FROM essays e
        LEFT JOIN assignment_submissions asub ON e.id = asub.essay_id
        LEFT JOIN assignments a ON asub.assignment_id = a.id
        LEFT JOIN users u ON a.teacher_id = u.id
        WHERE e.id = %s AND e.user_id = %s
    """, (essay_id, session['user_id']))
    essay = cursor.fetchone()
    
    conn.close()
    
    if not essay:
        flash('Essay not found', 'error')
        return redirect(url_for('student_dashboard'))
    
    if not essay[5]:  # No teacher feedback
        flash('No teacher feedback available for this essay', 'info')
        return redirect(url_for('student_dashboard'))
    
    return render_template('student/feedback.html', essay=essay)

@app.route('/analyze', methods=['POST'])
@login_required
def analyze_essay():
    try:
        data = request.get_json()
        
        if not data:
            logger.error("No JSON data received")
            return jsonify({'error': 'No data provided'}), 400
        
        # Support both 'essay_text' and 'essay' field names for backward compatibility
        essay_text = data.get('essay_text', data.get('essay', '')).strip()
        essay_type = data.get('essay_type', 'auto').lower()
        coaching_level = data.get('coaching_level', 'medium').lower()
        suggestion_aggressiveness = data.get('suggestion_aggressiveness', 'medium').lower()
        assignment_id = data.get('assignment_id')  # Optional assignment ID
        
        if not essay_text:
            return jsonify({'error': 'Essay text is required'}), 400
        
        # Validate essay length
        if len(essay_text) < 50:
            return jsonify({'error': 'Essay too short (minimum 50 characters)'}), 400
        if len(essay_text) > 10000:
            return jsonify({'error': 'Essay too long (maximum 10000 characters)'}), 400
            
        # Validate essay type
        valid_types = ['auto', 'argumentative', 'narrative', 'literary', 'expository', 'descriptive', 'compare']
        if essay_type not in valid_types:
            return jsonify({'error': f'Invalid essay type. Must be one of: {", ".join(valid_types)}'}), 400
            
        # Validate coaching level
        valid_levels = ['light', 'medium', 'intensive']
        if coaching_level not in valid_levels:
            return jsonify({'error': f'Invalid coaching level. Must be one of: {", ".join(valid_levels)}'}), 400

        # Validate suggestion aggressiveness
        valid_aggressiveness = ['low', 'medium', 'high']
        if suggestion_aggressiveness not in valid_aggressiveness:
            return jsonify({'error': f'Invalid suggestion aggressiveness. Must be one of: {", ".join(valid_aggressiveness)}'}), 400
        
        logger.info(f"Processing essay: {len(essay_text)} chars, type: {essay_type}, level: {coaching_level}, aggressiveness: {suggestion_aggressiveness}")
        
        # Analyze essay with AI
        analysis_result = analyze_essay_with_ai(
            essay_text, 
            essay_type, 
            coaching_level, 
            suggestion_aggressiveness
        )
        
        # Save to database if user wants to save
        if data.get('save_to_db', False):
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Save essay
            cursor.execute("""
                INSERT INTO essays (user_id, title, content, essay_type, feedback, 
                                  ideas_score, organization_score, style_score, grammar_score, 
                                  total_score, status, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                session['user_id'],
                data.get('title', 'Untitled Essay'),
                essay_text,
                analysis_result.get('essay_type', 'Unknown'),
                json.dumps(analysis_result),
                analysis_result.get('scores', {}).get('ideas', 0),
                analysis_result.get('scores', {}).get('organization', 0),
                analysis_result.get('scores', {}).get('style', 0),
                analysis_result.get('scores', {}).get('grammar', 0),
                sum(analysis_result.get('scores', {}).values()),
                'analyzed',
                datetime.datetime.now()
            ))
            conn.commit()
            conn.close()
        
        # Save submission to database for students
        user = get_current_user()
        if user and user['role'] == 'student' and analysis_result.get('analysis_id'):
            success = save_submission_to_db(user['id'], analysis_result.get('analysis_id'), assignment_id)
            if success:
                logger.info("Submission saved successfully")
            else:
                logger.warning("Failed to save submission to database")
        
        return jsonify(analysis_result)
        
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/export', methods=['POST'])
@login_required
def export_to_word():
    data = request.get_json()
    
    if not data or 'essay' not in data or 'analysis' not in data:
        return jsonify({'error': 'Missing essay or analysis data'}), 400
    
    try:
        essay_text = data['essay']
        analysis_data = data['analysis']
        accepted_suggestions = data.get('acceptedSuggestions', [])
        
        # Ensure analysis_data is a dict, parse if string
        if isinstance(analysis_data, str):
            analysis_data = json.loads(analysis_data)

        # Normalize suggestions to list of dicts
        suggestions = analysis_data.get('suggestions', [])
        normalized_suggestions = []
        for s in suggestions:
            if isinstance(s, dict):
                normalized_suggestions.append(s)
            elif isinstance(s, str):
                normalized_suggestions.append({'type': 'General', 'text': s, 'reason': ''})
            else:
                normalized_suggestions.append({'type': 'General', 'text': str(s), 'reason': ''})
        analysis_data['suggestions'] = normalized_suggestions

        # Debug logging
        logger.debug(f"Type of analysis_data: {type(analysis_data)}")
        if isinstance(analysis_data, dict):
            logger.debug(f"Keys in analysis_data: {list(analysis_data.keys())}")
            tagged_essay = analysis_data.get('tagged_essay', '')
            logger.debug(f"Tagged essay snippet: {tagged_essay[:100]}")
        else:
            logger.debug(f"analysis_data content snippet: {str(analysis_data)[:100]}")
        
        doc = create_word_document_with_suggestions(essay_text, analysis_data, accepted_suggestions)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        logger.info("Word document exported successfully")
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name='essay_analysis.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Export error: {e}")
        return jsonify({'error': f'Export failed: {str(e)}'}), 500

@app.route('/analytics')
@login_required
def get_analytics():
    connection = get_db_connection()
    if not connection:
        return jsonify({'error': 'Database connection failed'}), 500
    
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT essay_type, AVG(overall_score) as avg_score, COUNT(*) as count
                FROM essay_analyses 
                WHERE created_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)
                GROUP BY essay_type
            """)
            
            analytics_data = {'essay_types': [], 'avg_scores': [], 'counts': []}
            
            for row in cursor.fetchall():
                analytics_data['essay_types'].append(row[0])
                analytics_data['avg_scores'].append(float(row[1]))
                analytics_data['counts'].append(row[2])
            
            return jsonify(analytics_data)
    except Exception as e:
        logger.error(f"Analytics error: {e}")
        return jsonify({'error': f'Analytics failed: {str(e)}'}), 500
    finally:
        connection.close()

@app.route('/api/suggestions/accept', methods=['POST'])
@login_required
def accept_word_suggestion():
    """Accept a word-level suggestion"""
    try:
        data = request.get_json()
        suggestion_id = data.get('suggestion_id')
        suggestion_type = data.get('type')
        text = data.get('text')
        
        logger.info(f"User {session['user_id']} accepting suggestion: {suggestion_id}, type: {suggestion_type}, text: {text}")
        
        # In a real application, you might want to save this to database
        # For now, we'll just return success
        
        return jsonify({
            'success': True,
            'message': 'Suggestion accepted',
            'suggestion_id': suggestion_id,
            'type': suggestion_type,
            'action': 'accepted'
        })
        
    except Exception as e:
        logger.error(f"Error accepting suggestion: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/suggestions/reject', methods=['POST'])
@login_required
def reject_word_suggestion():
    """Reject a word-level suggestion"""
    try:
        data = request.get_json()
        suggestion_id = data.get('suggestion_id')
        suggestion_type = data.get('type')
        text = data.get('text')
        
        logger.info(f"User {session['user_id']} rejecting suggestion: {suggestion_id}, type: {suggestion_type}, text: {text}")
        
        # In a real application, you might want to save this to database
        # For now, we'll just return success
        
        return jsonify({
            'success': True,
            'message': 'Suggestion rejected',
            'suggestion_id': suggestion_id,
            'type': suggestion_type,
            'action': 'rejected'
        })
        
    except Exception as e:
        logger.error(f"Error rejecting suggestion: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/checklist/<int:analysis_id>/progress', methods=['GET'])
@login_required
def get_checklist_progress_api(analysis_id):
    """Get checklist progress for an analysis"""
    try:
        steps = get_checklist_progress(analysis_id)
        return jsonify({
            'success': True,
            'steps': steps
        })
    except Exception as e:
        logger.error(f"Error getting checklist progress: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/checklist/<int:analysis_id>/update', methods=['POST'])
@login_required
def update_checklist_progress_api(analysis_id):
    """Update checklist step progress"""
    try:
        data = request.get_json()
        step_name = data.get('step_name')
        completed = data.get('completed', True)
        
        if not step_name:
            return jsonify({'error': 'Step name is required'}), 400
        
        success = update_checklist_progress(analysis_id, step_name, completed)
        
        if success:
            # Get updated progress
            updated_steps = get_checklist_progress(analysis_id)
            return jsonify({
                'success': True,
                'message': 'Checklist updated successfully',
                'steps': updated_steps
            })
        else:
            return jsonify({'error': 'Failed to update checklist'}), 500
            
    except Exception as e:
        logger.error(f"Error updating checklist progress: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/rubric/configure', methods=['POST'])
@login_required
@role_required('teacher')
def configure_custom_rubric():
    """Configure a custom rubric for a teacher"""
    try:
        data = request.get_json()
        rubric_name = data.get('name')
        weights = data.get('weights', {})
        custom_criteria = data.get('custom_criteria')
        
        # Validate weights sum to 100
        total_weight = sum(weights.values())
        if total_weight != 100:
            return jsonify({'error': 'Rubric weights must sum to 100'}), 400
        
        rubric_id = create_modular_rubric_engine(
            session['user_id'], 
            rubric_name, 
            weights, 
            custom_criteria
        )
        
        if rubric_id:
            return jsonify({
                'success': True,
                'rubric_id': rubric_id,
                'message': 'Custom rubric created successfully'
            })
        else:
            return jsonify({'error': 'Failed to create rubric'}), 500
            
    except Exception as e:
        logger.error(f"Error configuring rubric: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics/student/<int:student_id>')
@login_required
@role_required('teacher')
def get_student_analytics(student_id):
    """Get detailed analytics for a specific student"""
    try:
        connection = get_db_connection()
        if not connection:
            return jsonify({'error': 'Database connection failed'}), 500
        
        with connection.cursor() as cursor:
            # Verify teacher has access to this student
            cursor.execute("""
                SELECT id FROM student_teacher_assignments 
                WHERE student_id = %s AND teacher_id = %s
            """, (student_id, session['user_id']))
            
            if not cursor.fetchone():
                return jsonify({'error': 'Access denied'}), 403
            
            # Get progress over time
            cursor.execute("""
                SELECT DATE(ea.created_at) as date, 
                       AVG(ea.overall_score) as avg_score,
                       COUNT(*) as essay_count
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s
                GROUP BY DATE(ea.created_at)
                ORDER BY date DESC
                LIMIT 30
            """, (student_id,))
            progress_data = cursor.fetchall()
            
            # Get rubric dimension trends
            cursor.execute("""
                SELECT 'ideas' as dimension, AVG(ea.ideas_score) as avg_score
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s
                UNION ALL
                SELECT 'organization', AVG(ea.organization_score)
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s
                UNION ALL
                SELECT 'style', AVG(ea.style_score)
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s
                UNION ALL
                SELECT 'grammar', AVG(ea.grammar_score)
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s
            """, (student_id, student_id, student_id, student_id))
            dimension_scores = cursor.fetchall()
            
            # Check for at-risk indicators
            cursor.execute("""
                SELECT AVG(ea.overall_score) as recent_avg
                FROM essay_analyses ea
                JOIN student_submissions ss ON ea.id = ss.analysis_id
                WHERE ss.student_id = %s 
                AND ea.created_at >= DATE_SUB(NOW(), INTERVAL 14 DAY)
            """, (student_id,))
            recent_avg = cursor.fetchone()[0] or 0
            
            at_risk = recent_avg < 60  # Below 60% average in last 2 weeks
            
            analytics_data = {
                'progress_data': [
                    {
                        'date': row[0].isoformat() if row[0] else None,
                        'avg_score': float(row[1]) if row[1] else 0,
                        'essay_count': row[2]
                    } for row in progress_data
                ],
                'dimension_scores': {
                    row[0]: float(row[1]) if row[1] else 0 
                    for row in dimension_scores
                },
                'at_risk': at_risk,
                'recent_average': float(recent_avg) if recent_avg else 0
            }
            
            return jsonify(analytics_data)
            
    except Exception as e:
        logger.error(f"Error getting student analytics: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        if 'connection' in locals():
            connection.close()

@app.route('/student/settings', methods=['GET', 'POST'])
@login_required
@role_required('student')
def student_settings():
    """Student settings page"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        # Handle settings update
        email = request.form.get('email')
        default_essay_type = request.form.get('default_essay_type', 'auto')
        coaching_level = request.form.get('coaching_level', 'medium')
        suggestion_aggressiveness = request.form.get('suggestion_aggressiveness', 'medium')
        
        # Update user settings
        cursor.execute("""
            UPDATE users SET 
                email = %s, 
                default_essay_type = %s, 
                coaching_level = %s, 
                suggestion_aggressiveness = %s 
            WHERE id = %s
        """, (email, default_essay_type, coaching_level, suggestion_aggressiveness, session['user_id']))
        conn.commit()
        
        flash('Settings updated successfully!', 'success')
        return redirect(url_for('student_settings'))
    
    # Get current user settings
    cursor.execute("""
        SELECT email, default_essay_type, coaching_level, suggestion_aggressiveness
        FROM users WHERE id = %s
    """, (session['user_id'],))
    user_settings = cursor.fetchone()
    conn.close()
    
    if not user_settings:
        flash('User not found', 'error')
        return redirect(url_for('student_dashboard'))
    
    return render_template('student/settings.html', user_settings=user_settings)

@app.route('/teacher/settings', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def teacher_settings():
    """Teacher settings page"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        # Handle settings update
        email = request.form.get('email')
        default_feedback_template = request.form.get('default_feedback_template', '')
        auto_grade_essays = request.form.get('auto_grade_essays') == 'on'
        
        # Update user settings
        cursor.execute("""
            UPDATE users SET 
                email = %s, 
                default_feedback_template = %s, 
                auto_grade_essays = %s 
            WHERE id = %s
        """, (email, default_feedback_template, auto_grade_essays, session['user_id']))
        conn.commit()
        
        flash('Settings updated successfully!', 'success')
        return redirect(url_for('teacher_settings'))
    
    # Get current user settings
    cursor.execute("""
        SELECT email, default_feedback_template, auto_grade_essays
        FROM users WHERE id = %s
    """, (session['user_id'],))
    user_settings = cursor.fetchone()
    conn.close()
    
    if not user_settings:
        flash('User not found', 'error')
        return redirect(url_for('teacher_dashboard'))
    
    return render_template('teacher/settings.html', user_settings=user_settings)

# Custom Jinja2 filters
@app.template_filter('nl2br')
def nl2br_filter(text):
    """Convert newlines to HTML breaks"""
    return text.replace('\n', '<br>') if text else ''

# Additional missing routes
@app.route('/teacher/analytics')
@login_required
@role_required('teacher')
def teacher_analytics():
    """Teacher analytics page"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get analytics data for assigned students
    cursor.execute("""
        SELECT 
            COUNT(DISTINCT e.id) as total_essays,
            AVG(e.total_score) as avg_score,
            COUNT(DISTINCT sta.student_id) as total_students
        FROM student_teacher_assignments sta
        LEFT JOIN essays e ON sta.student_id = e.user_id
        WHERE sta.teacher_id = %s
    """, (session['user_id'],))
    overall_stats = cursor.fetchone()
    
    # Get score trends over time
    cursor.execute("""
        SELECT DATE(e.created_at) as date, AVG(e.total_score) as avg_score
        FROM essays e
        JOIN student_teacher_assignments sta ON e.user_id = sta.student_id
        WHERE sta.teacher_id = %s AND e.total_score IS NOT NULL
        GROUP BY DATE(e.created_at)
        ORDER BY date DESC
        LIMIT 30
    """, (session['user_id'],))
    score_trends = cursor.fetchall()
    
    # Get essay type distribution
    cursor.execute("""
        SELECT e.essay_type, COUNT(*) as count
        FROM essays e
        JOIN student_teacher_assignments sta ON e.user_id = sta.student_id
        WHERE sta.teacher_id = %s
        GROUP BY e.essay_type
    """, (session['user_id'],))
    essay_types = cursor.fetchall()
    
    conn.close()
    
    return render_template('teacher/analytics.html', 
                         overall_stats=overall_stats,
                         score_trends=score_trends,
                         essay_types=essay_types)

@app.route('/student/profile', methods=['GET', 'POST'])
@login_required
@role_required('student')
def student_profile():
    """Student profile page"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        # Handle profile update
        username = request.form.get('username')
        email = request.form.get('email')
        
        # Update user profile
        cursor.execute("""
            UPDATE users SET username = %s, email = %s WHERE id = %s
        """, (username, email, session['user_id']))
        conn.commit()
        
        # Update session
        session['username'] = username
        
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('student_profile'))
    
    # Get current user info
    cursor.execute("""
        SELECT username, email, created_at
        FROM users WHERE id = %s
    """, (session['user_id'],))
    user_info = cursor.fetchone()
    
    conn.close()
    
    return render_template('student/profile.html', user_info=user_info)

@app.route('/teacher/profile', methods=['GET', 'POST'])
@login_required
@role_required('teacher')
def teacher_profile():
    """Teacher profile page"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        # Handle profile update
        username = request.form.get('username')
        email = request.form.get('email')
        
        # Update user profile
        cursor.execute("""
            UPDATE users SET username = %s, email = %s WHERE id = %s
        """, (username, email, session['user_id']))
        conn.commit()
        
        # Update session
        session['username'] = username
        
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('teacher_profile'))
    
    # Get current user info
    cursor.execute("""
        SELECT username, email, created_at
        FROM users WHERE id = %s
    """, (session['user_id'],))
    user_info = cursor.fetchone()
    
    conn.close()
    
    return render_template('teacher/profile.html', user_info=user_info)

@app.route('/api/essays/recent')
@login_required
def get_recent_essays():
    """Get recent essays for dashboard"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if session['role'] == 'student':
        cursor.execute("""
            SELECT id, title, essay_type, total_score, created_at, status
            FROM essays 
            WHERE user_id = %s 
            ORDER BY created_at DESC 
            LIMIT 5
        """, (session['user_id'],))
    else:
        cursor.execute("""
            SELECT e.id, e.title, e.essay_type, e.total_score, e.created_at, e.status, u.username
            FROM essays e
            JOIN users u ON e.user_id = u.id
            JOIN student_teacher_assignments sta ON u.id = sta.student_id
            WHERE sta.teacher_id = %s
            ORDER BY e.created_at DESC
            LIMIT 10
        """, (session['user_id'],))
    
    essays = cursor.fetchall()
    conn.close()
    
    # Convert to list of dicts for JSON response
    essay_list = []
    for essay in essays:
        essay_dict = {
            'id': essay[0],
            'title': essay[1],
            'essay_type': essay[2],
            'total_score': essay[3],
            'created_at': essay[4].isoformat() if essay[4] else None,
            'status': essay[5]
        }
        if len(essay) > 6:  # Teacher view includes username
            essay_dict['username'] = essay[6]
        essay_list.append(essay_dict)
    
    return jsonify(essay_list)

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500

if __name__ == '__main__':
    app.run(host='10.10.12.31', port=5000, debug=True)
